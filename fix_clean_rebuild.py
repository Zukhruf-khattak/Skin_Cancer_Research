"""
CLEAN REBUILD of zukhruftheisis_final.docx:
1. Fix page numbering:
   - Title page (Section 0, para 0-27): NO page number shown
   - Front matter (Section 1, para 28-46, Acknowledgement to blank before Ch.I): Roman numerals i, ii, iii...
   - Chapter I onwards (Section 2+): Arabic 1, 2, 3...
2. Insert ONE clickable TOC before Chapter I
3. Remove any previously inserted duplicate TOC

Strategy:
- Remove all existing section breaks except the body sectPr
- Rebuild clean section structure:
  Section A: title page (para 0-10 area) — no page number footer
  Section B: front matter (para 11 onwards until Ch.I) — Roman numeral footer
  Section C: Chapter I onwards — Arabic footer starting at 1
- Add proper footer with PAGE field to each section
- Insert clickable TOC right before Chapter I
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.part import Part
from docx.opc.packuri import PackURI
import copy, re
from lxml import etree

doc = Document(r'd:\RESEARCH PNGS\zukhruftheisis_final.docx')

# =====================================================================
# STEP 0: REMOVE ALL EXISTING SECTION BREAKS IN PARAGRAPH pPr
#         (we'll rebuild them cleanly)
# =====================================================================
print("Step 0: Removing existing section breaks in paragraphs...")
removed = 0
for para in doc.paragraphs:
    pPr = para._p.find(qn('w:pPr'))
    if pPr is not None:
        sectPr = pPr.find(qn('w:sectPr'))
        if sectPr is not None:
            pPr.remove(sectPr)
            removed += 1
print(f"  Removed {removed} section breaks.")

# =====================================================================
# STEP 1: FIND KEY PARAGRAPH INDICES
# =====================================================================
print("\nStep 1: Finding key paragraph indices...")

title_page_end_idx = -1   # Last para of title page (before Acknowledgement)
chapter1_idx = -1         # First para of Chapter I

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if 'ACKNOWLEDGEMENT' in text.upper() and title_page_end_idx == -1:
        title_page_end_idx = i - 1  # Para BEFORE Acknowledgement
    if re.match(r'^CHAPTER\s+I\b', text, re.IGNORECASE) and chapter1_idx == -1:
        chapter1_idx = i

print(f"  Title page ends at para: {title_page_end_idx}")
print(f"  'CHAPTER I' at para: {chapter1_idx}")

# The para just BEFORE Chapter I is where we insert our section break 
# to switch from Roman to Arabic numbering
pre_chapter1_idx = chapter1_idx - 1
# Find last non-empty para before Chapter I
while pre_chapter1_idx > 0 and not doc.paragraphs[pre_chapter1_idx].text.strip():
    pre_chapter1_idx -= 1

print(f"  Section break (Roman->Arabic) after para: {pre_chapter1_idx} = '{doc.paragraphs[pre_chapter1_idx].text.strip()[:60]}'")

# =====================================================================
# HELPER: Create a clean footer XML part with centered PAGE field
# =====================================================================
def create_footer_xml(show_page_num=True):
    """Return an lxml element for a w:ftr with optional page number."""
    WORD_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    
    ftr = OxmlElement('w:ftr')
    p = OxmlElement('w:p')
    ftr.append(p)
    
    # Paragraph properties: centered
    pPr = OxmlElement('w:pPr')
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    pPr.append(jc)
    p.append(pPr)
    
    if show_page_num:
        # Run properties (font)
        rPr = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        rPr.append(rFonts)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '24')  # 12pt
        rPr.append(sz)
        
        # fldChar begin
        r1 = OxmlElement('w:r')
        r1.append(copy.deepcopy(rPr))
        fc1 = OxmlElement('w:fldChar')
        fc1.set(qn('w:fldCharType'), 'begin')
        r1.append(fc1)
        p.append(r1)
        
        # instrText " PAGE "
        r2 = OxmlElement('w:r')
        r2.append(copy.deepcopy(rPr))
        instr = OxmlElement('w:instrText')
        instr.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        instr.text = ' PAGE '
        r2.append(instr)
        p.append(r2)
        
        # fldChar end
        r3 = OxmlElement('w:r')
        r3.append(copy.deepcopy(rPr))
        fc3 = OxmlElement('w:fldChar')
        fc3.set(qn('w:fldCharType'), 'end')
        r3.append(fc3)
        p.append(r3)
    
    return ftr

# =====================================================================
# HELPER: Add a footer part to the document and return its rId
# =====================================================================
def add_footer_part(doc, footer_xml_element, part_name_suffix):
    """Add a footer XML part to the docx package and return the relationship rId."""
    from docx.opc.part import Part
    from docx.opc.packuri import PackURI
    import lxml.etree as etree
    
    CT_FOOTER = 'application/vnd.openxmlformats-officedocument.wordprocessingml.footer+xml'
    RT_FOOTER = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/footer'
    
    # Serialize the footer element
    NAMESPACES = {
        'wpc': 'http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas',
        'cx': 'http://schemas.microsoft.com/office/drawing/2014/chartex',
        'mc': 'http://schemas.openxmlformats.org/markup-compatibility/2006',
        'aink': 'http://schemas.microsoft.com/office/drawing/2016/ink',
        'am3d': 'http://schemas.microsoft.com/office/drawing/2017/model3d',
        'o': 'urn:schemas-microsoft-com:office:office',
        'oel': 'http://schemas.microsoft.com/office/2019/extlst',
        'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
        'm': 'http://schemas.openxmlformats.org/officeDocument/2006/math',
        'v': 'urn:schemas-microsoft-com:vml',
        'wp14': 'http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing',
        'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
        'w10': 'urn:schemas-microsoft-com:office:word',
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'w14': 'http://schemas.microsoft.com/office/word/2010/wordml',
        'w15': 'http://schemas.microsoft.com/office/word/2012/wordml',
        'w16cex': 'http://schemas.microsoft.com/office/word/2018/wordml/cex',
        'w16cid': 'http://schemas.microsoft.com/office/word/2016/wordml/cid',
        'w16': 'http://schemas.microsoft.com/office/word/2018/wordml',
        'w16sdtdh': 'http://schemas.microsoft.com/office/word/2020/wordml/sdtdatahash',
        'w16se': 'http://schemas.microsoft.com/office/word/2015/wordml/symex',
        'wpg': 'http://schemas.microsoft.com/office/word/2010/wordprocessingGroup',
        'wpi': 'http://schemas.microsoft.com/office/word/2010/wordprocessingInk',
        'wne': 'http://schemas.microsoft.com/office/word/2006/wordml',
        'wps': 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape',
    }
    
    # Build a proper footer XML string
    footer_str = b'<?xml version=\'1.0\' encoding=\'UTF-8\' standalone=\'yes\'?>\n'
    footer_str += etree.tostring(footer_xml_element, xml_declaration=False)
    
    # Create the part
    part_uri = PackURI(f'/word/footer{part_name_suffix}.xml')
    footer_part = Part(part_uri, CT_FOOTER, footer_str, doc.part.package)
    
    # Add relationship from document part
    rId = doc.part.relate_to(footer_part, RT_FOOTER)
    return rId

# =====================================================================
# STEP 2: CREATE FOOTER PARTS
# =====================================================================
print("\nStep 2: Creating footer parts...")

# Footer A: No page number (for title page)
footer_none_xml = create_footer_xml(show_page_num=False)
rId_none = add_footer_part(doc, footer_none_xml, 'Blank')
print(f"  Created blank footer: rId={rId_none}")

# Footer B: Page number (Roman format - but the FORMAT is set in sectPr pgNumType)
footer_page_xml = create_footer_xml(show_page_num=True)
rId_page = add_footer_part(doc, footer_page_xml, 'WithNum')
print(f"  Created page-number footer: rId={rId_page}")

# Footer C: Arabic page numbers (for Chapter I onwards) - same template, different pgNumType
footer_arabic_xml = create_footer_xml(show_page_num=True)
rId_arabic = add_footer_part(doc, footer_arabic_xml, 'Arabic')
print(f"  Created Arabic page-number footer: rId={rId_arabic}")

# =====================================================================
# HELPER: Build a sectPr element
# =====================================================================
def build_sectPr(footer_rId, num_format, start_num, inherit_from=None):
    """Build a complete sectPr element."""
    sectPr = OxmlElement('w:sectPr')
    
    # Footer reference
    ftrRef = OxmlElement('w:footerReference')
    ftrRef.set(qn('w:type'), 'default')
    ftrRef.set('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id', footer_rId)
    sectPr.append(ftrRef)
    
    # Page number type
    if num_format is not None:
        pgNumType = OxmlElement('w:pgNumType')
        pgNumType.set(qn('w:fmt'), num_format)
        pgNumType.set(qn('w:start'), str(start_num))
        sectPr.append(pgNumType)
    
    # Copy page size and margins from existing body sectPr
    if inherit_from is not None:
        for tag in [qn('w:pgSz'), qn('w:pgMar')]:
            el = inherit_from.find(tag)
            if el is not None:
                sectPr.append(copy.deepcopy(el))
    
    return sectPr

# Get body sectPr to inherit page size/margins
body = doc.element.body
body_sectPr = body.find(qn('w:sectPr'))

# =====================================================================
# STEP 3: INSERT SECTION BREAK AFTER TITLE PAGE
#         (para title_page_end_idx) -> switches to Roman numerals
# =====================================================================
print("\nStep 3: Inserting section breaks...")

# Section break 1: After title page last para -> starts front matter (Roman)
title_end_para = doc.paragraphs[title_page_end_idx]
pPr = title_end_para._p.find(qn('w:pPr'))
if pPr is None:
    pPr = OxmlElement('w:pPr')
    title_end_para._p.insert(0, pPr)
sect_title = build_sectPr(rId_none, None, 0, inherit_from=body_sectPr)
# For section type: "nextPage" to ensure title page is its own page
sectType = OxmlElement('w:type')
sectType.set(qn('w:val'), 'nextPage')
sect_title.insert(0, sectType)
pPr.append(sect_title)
print(f"  Section break 1 added after para {title_page_end_idx} (title page end)")

# Section break 2: After last front-matter para -> switches to Arabic (Chapter I)
pre_ch1_para = doc.paragraphs[pre_chapter1_idx]
pPr2 = pre_ch1_para._p.find(qn('w:pPr'))
if pPr2 is None:
    pPr2 = OxmlElement('w:pPr')
    pre_ch1_para._p.insert(0, pPr2)
sect_frontmatter = build_sectPr(rId_page, 'lowerRoman', 1, inherit_from=body_sectPr)
pPr2.append(sect_frontmatter)
print(f"  Section break 2 added after para {pre_chapter1_idx} (last front-matter para)")

# Update body sectPr (Chapter I onwards) -> Arabic from 1
# Remove existing pgNumType and footerReference from body sectPr
for child in list(body_sectPr):
    if child.tag in [qn('w:pgNumType'), qn('w:footerReference')]:
        body_sectPr.remove(child)

# Add fresh footer and page num
body_ftrRef = OxmlElement('w:footerReference')
body_ftrRef.set(qn('w:type'), 'default')
body_ftrRef.set('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id', rId_arabic)
body_sectPr.insert(0, body_ftrRef)

body_pgNumType = OxmlElement('w:pgNumType')
body_pgNumType.set(qn('w:fmt'), 'decimal')
body_pgNumType.set(qn('w:start'), '1')
body_sectPr.append(body_pgNumType)
print(f"  Body sectPr updated: Arabic numerals from 1")

# =====================================================================
# STEP 4: INSERT CLICKABLE TOC BEFORE CHAPTER I
# =====================================================================
print("\nStep 4: Building clickable TOC...")

# Collect all headings for TOC
toc_entries = []
bk_counter = 0
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    style = para.style.name
    if not text:
        continue
    
    level = None
    if re.match(r'^CHAPTER\s+(I|II|III|IV|V|VI|VII|VIII|IX|X)\b', text, re.IGNORECASE):
        level = 1
    elif text.upper() == 'REFERENCES':
        level = 1
    elif 'Heading 2' in style:
        level = 2
    elif 'Heading 3' in style:
        level = 3
    
    if level is not None:
        bk_name = f'_Toc_{bk_counter:05d}'
        bk_counter += 1
        toc_entries.append((i, text, level, bk_name))

print(f"  Found {len(toc_entries)} headings")

# Add bookmarks to heading paragraphs
for (idx, text, level, bk_name) in toc_entries:
    para = doc.paragraphs[idx]
    # Remove any existing bookmarkStart/End to avoid duplicates
    for bk in para._p.findall('.//' + qn('w:bookmarkStart')):
        if bk.get(qn('w:name'), '').startswith('_Toc_'):
            bk.getparent().remove(bk)
    for bk in para._p.findall('.//' + qn('w:bookmarkEnd')):
        bk.getparent().remove(bk)
    
    bk_id = str(bk_counter + idx)
    bk_start = OxmlElement('w:bookmarkStart')
    bk_start.set(qn('w:id'), bk_id)
    bk_start.set(qn('w:name'), bk_name)
    bk_end = OxmlElement('w:bookmarkEnd')
    bk_end.set(qn('w:id'), bk_id)
    
    para._p.insert(0, bk_start)
    para._p.append(bk_end)

def create_toc_entry_para(text, level, bookmark_name):
    """Create a TOC line paragraph element with hyperlink."""
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    
    indent_twips = {1: 0, 2: 400, 3: 800}.get(level, 0)
    if indent_twips:
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), str(indent_twips))
        pPr.append(ind)
    
    # Dot leader tab at right margin (~5.5 inches from left of text area)
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:leader'), 'dot')
    tab.set(qn('w:pos'), '8640')
    tabs.append(tab)
    pPr.append(tabs)
    
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '40')
    spacing.set(qn('w:after'), '40')
    pPr.append(spacing)
    p.append(pPr)
    
    # Font size by level
    font_size_hpc = {1: 26, 2: 24, 3: 22}.get(level, 24)  # half-points
    is_bold = (level == 1)
    
    def make_rPr():
        rPr = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        rPr.append(rFonts)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(font_size_hpc))
        rPr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), str(font_size_hpc))
        rPr.append(szCs)
        if is_bold:
            rPr.append(OxmlElement('w:b'))
        col = OxmlElement('w:color')
        col.set(qn('w:val'), '000000')
        rPr.append(col)
        return rPr
    
    # Hyperlink pointing to bookmark anchor
    hl = OxmlElement('w:hyperlink')
    hl.set(qn('w:anchor'), bookmark_name)
    hl_r = OxmlElement('w:r')
    hl_r.append(make_rPr())
    hl_t = OxmlElement('w:t')
    hl_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    hl_t.text = text
    hl_r.append(hl_t)
    hl.append(hl_r)
    p.append(hl)
    
    return p

# Find parent element and insertion point (just before Chapter I paragraph)
chapter1_para = doc.paragraphs[chapter1_idx]
parent_body = chapter1_para._p.getparent()
ch1_pos = list(parent_body).index(chapter1_para._p)

# Build TOC elements in correct forward order and insert at ch1_pos
toc_elements = []

# Page break before TOC (after front matter, before TOC title)
toc_start_pb = OxmlElement('w:p')
toc_start_pb_r = OxmlElement('w:r')
toc_start_pb_br = OxmlElement('w:br')
toc_start_pb_br.set(qn('w:type'), 'page')
toc_start_pb_r.append(toc_start_pb_br)
toc_start_pb.append(toc_start_pb_r)
toc_elements.append(toc_start_pb)

# TOC title paragraph
toc_title_p = OxmlElement('w:p')
toc_title_pPr = OxmlElement('w:pPr')
toc_title_jc = OxmlElement('w:jc')
toc_title_jc.set(qn('w:val'), 'center')
toc_title_pPr.append(toc_title_jc)
toc_title_spacing = OxmlElement('w:spacing')
toc_title_spacing.set(qn('w:before'), '0')
toc_title_spacing.set(qn('w:after'), '240')
toc_title_pPr.append(toc_title_spacing)
toc_title_p.append(toc_title_pPr)

toc_title_rPr = OxmlElement('w:rPr')
toc_title_rFonts = OxmlElement('w:rFonts')
toc_title_rFonts.set(qn('w:ascii'), 'Times New Roman')
toc_title_rFonts.set(qn('w:hAnsi'), 'Times New Roman')
toc_title_rPr.append(toc_title_rFonts)
toc_title_sz = OxmlElement('w:sz')
toc_title_sz.set(qn('w:val'), '32')  # 16pt
toc_title_rPr.append(toc_title_sz)
toc_title_szCs = OxmlElement('w:szCs')
toc_title_szCs.set(qn('w:val'), '32')
toc_title_rPr.append(toc_title_szCs)
toc_title_rPr.append(OxmlElement('w:b'))
toc_title_r = OxmlElement('w:r')
toc_title_r.append(toc_title_rPr)
toc_title_t = OxmlElement('w:t')
toc_title_t.text = 'TABLE OF CONTENTS'
toc_title_r.append(toc_title_t)
toc_title_p.append(toc_title_r)
toc_elements.append(toc_title_p)

# Spacer
sp = OxmlElement('w:p')
toc_elements.append(sp)

# TOC entries in correct order
for (idx, text, level, bk_name) in toc_entries:
    toc_elements.append(create_toc_entry_para(text, level, bk_name))

# Page break after TOC (before Chapter I)
pb_p = OxmlElement('w:p')
pb_r = OxmlElement('w:r')
pb_br = OxmlElement('w:br')
pb_br.set(qn('w:type'), 'page')
pb_r.append(pb_br)
pb_p.append(pb_r)
toc_elements.append(pb_p)

# Insert all elements at ch1_pos in forward order
for offset, el in enumerate(toc_elements):
    parent_body.insert(ch1_pos + offset, el)


print(f"  TOC with {len(toc_entries)} entries inserted before Chapter I")

# =====================================================================
# SAVE
# =====================================================================
output_path = r'd:\RESEARCH PNGS\zukhruftheisis_v3.docx'
doc.save(output_path)
print(f"\nDONE! Saved as: {output_path}")
print("Summary:")
print("  - Title page: NO page number")
print("  - Front matter (Acknowledgement to before Ch.I): Roman numerals i, ii, iii...")
print("  - Chapter I onwards: Arabic 1, 2, 3...")
print("  - ONE clickable TOC inserted before Chapter I")

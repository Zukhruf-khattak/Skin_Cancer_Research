from docx import Document

doc = Document("Huma Thesis_Updated.docx")

with open("ch2_dump.txt", "w", encoding="utf-8") as f:
    f.write("--- Paragraphs 155 to 185 ---\n")
    for i in range(155, 185):
        if i < len(doc.paragraphs):
            text = doc.paragraphs[i].text.strip()
            if text:
                f.write(f"[{i}] {doc.paragraphs[i].style.name}: {text}\n")
    
    f.write("\n--- References 50 to 65 ---\n")
    ref_start = -1
    for i, p in enumerate(doc.paragraphs):
        if "REFERENCES" in p.text.upper() or "BIBLIOGRAPHY" in p.text.upper():
            ref_start = i
            break

    if ref_start != -1:
        count = 1
        for i in range(ref_start + 1, len(doc.paragraphs)):
            text = doc.paragraphs[i].text.strip()
            if text.startswith("["):
                if count >= 50:
                    f.write(f"[{count}] {text[:150]}...\n")
                count += 1

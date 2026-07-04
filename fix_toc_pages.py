import docx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def fix_toc_pages(in_file, out_file):
    doc = docx.Document(in_file)
    
    # 1. Map paragraphs to bookmarks
    bkmk_to_para = {}
    for i, p in enumerate(doc.paragraphs):
        for bk in p._p.findall('.//' + qn('w:bookmarkStart')):
            name = bk.get(qn('w:name'), '')
            if name.startswith('TOC_BM_'):
                bkmk_to_para[name] = i
                
    if 'TOC_BM_0' not in bkmk_to_para:
        print("TOC_BM_0 not found!")
        return
        
    start_idx = bkmk_to_para['TOC_BM_0']
    print(f"Chapter 1 starts at physical paragraph: {start_idx}")
    
    # 2. Count pages starting from start_idx
    current_page = 1
    mapping = {}
    
    for i in range(start_idx, len(doc.paragraphs)):
        p = doc.paragraphs[i]
        
        # Check if this paragraph has a heading bookmark
        for bk in p._p.findall('.//' + qn('w:bookmarkStart')):
            name = bk.get(qn('w:name'), '')
            if name.startswith('TOC_BM_'):
                mapping[name] = str(current_page)
                
        # Count breaks for the next paragraph
        breaks = p._p.findall('.//' + qn('w:br'))
        rendered = p._p.findall('.//' + qn('w:lastRenderedPageBreak'))
        num_breaks = sum(1 for br in breaks if br.get(qn('w:type')) == 'page')
        num_breaks += len(rendered)
        
        # Section breaks often act as page breaks
        pPr = p._p.find(qn('w:pPr'))
        if pPr is not None and pPr.find(qn('w:sectPr')) is not None:
            num_breaks += 1
            
        current_page += num_breaks

    print(f"Total calculated pages for body: {current_page}")
    print(f"Sample mapping: TOC_BM_0 -> {mapping.get('TOC_BM_0')}, TOC_BM_1 -> {mapping.get('TOC_BM_1')}")
    print(f"Last mapping: TOC_BM_101 -> {mapping.get('TOC_BM_101')}")
    
    # 3. Update the TOC IN PLACE
    updated = 0
    for i, p in enumerate(doc.paragraphs):
        # Find hyperlinks that point to TOC bookmarks
        hls = p._p.findall('.//' + qn('w:hyperlink'))
        for hl in hls:
            anchor = hl.get(qn('w:anchor'), '')
            if anchor.startswith('TOC_BM_'):
                # We found a TOC entry
                page_str = mapping.get(anchor, "1")
                
                # The page number was added as a text node at the end of the paragraph.
                # Let's find the last w:t inside the last w:r in the paragraph and update it.
                # Wait, in the LAST script, we appended them!
                # But we are reading zukhruftheisis_FINAL_TOC.docx, which does NOT have the numbers yet!
                # Wait, zukhruftheisis_FINAL_TOC.docx didn't have page numbers. 
                # Oh, wait! The user said "you still didnt mention corectly because chapter 1 start from page 1 not 3... in zukhruftheisis_FINAL_TOC"
                # The user probably looked at zukhruftheisis_FINAL_TOC_Numbered.docx!
                # Wait, their prompt specifically says: "you can check and now mention page number correctly(in toc) in zukhruftheisis_FINAL_TOC"
                # They probably want me to re-run the injection into zukhruftheisis_FINAL_TOC.docx with the CORRECT page numbers.
                
                # Check if it already has tabs (meaning we injected it already)
                existing_tabs = p._p.findall('.//' + qn('w:tab'))
                if len(existing_tabs) == 0:
                    # Inject from scratch
                    pPr = p._p.find(qn('w:pPr'))
                    if pPr is None:
                        pPr = OxmlElement('w:pPr')
                        p._p.insert(0, pPr)
                    tabs = pPr.find(qn('w:tabs'))
                    if tabs is None:
                        tabs = OxmlElement('w:tabs')
                        pPr.append(tabs)
                    tab_stop = OxmlElement('w:tab')
                    tab_stop.set(qn('w:val'), 'right')
                    tab_stop.set(qn('w:leader'), 'dot')
                    tab_stop.set(qn('w:pos'), '9000') 
                    tabs.append(tab_stop)
                    
                    r = OxmlElement('w:r')
                    rPr_r = OxmlElement('w:rPr')
                    rFonts = OxmlElement('w:rFonts')
                    rFonts.set(qn('w:ascii'), 'Times New Roman')
                    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                    rPr_r.append(rFonts)
                    sz = OxmlElement('w:sz')
                    
                    # Get font size based on level
                    # If it's a chapter, level is 1
                    level = 2
                    text_run = hl.find('.//' + qn('w:t'))
                    if text_run is not None and text_run.text and text_run.text.strip().upper().startswith("CHAPTER"):
                        level = 1
                    elif text_run is not None and text_run.text and text_run.text.strip().upper() == "REFERENCES":
                        level = 1
                        
                    font_size = {1: 24, 2: 24, 3: 24}.get(level, 24)
                    sz.set(qn('w:val'), str(font_size))
                    rPr_r.append(sz)
                    
                    col = OxmlElement('w:color')
                    col.set(qn('w:val'), '000000')
                    rPr_r.append(col)
                    
                    if level == 1:
                        rPr_r.append(OxmlElement('w:b'))
                    
                    r.append(rPr_r)
                    r.append(OxmlElement('w:tab'))
                    
                    t = OxmlElement('w:t')
                    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                    t.text = page_str
                    r.append(t)
                    
                    p._p.append(r)
                    updated += 1
                else:
                    # It already has the page number from a previous run, just update the last text node
                    runs = p._p.findall('.//' + qn('w:r'))
                    if runs:
                        last_r = runs[-1]
                        t_nodes = last_r.findall('.//' + qn('w:t'))
                        if t_nodes:
                            t_nodes[-1].text = page_str
                            updated += 1

    print(f"Updated {updated} TOC entries!")
    doc.save(out_file)
    print(f"Saved to {out_file}")

fix_toc_pages(r'd:\RESEARCH PNGS\zukhruftheisis_FINAL_TOC.docx', r'd:\RESEARCH PNGS\zukhruftheisis_FINAL_TOC_Numbered.docx')

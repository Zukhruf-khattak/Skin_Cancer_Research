import docx
from docx.oxml.ns import qn

def check_footers(filename):
    doc = docx.Document(filename)
    print(f"=== Checking footers in {filename} ===")
    
    for i, section in enumerate(doc.sections):
        print(f"\nSection {i}:")
        sectPr = section._sectPr
        pnt = sectPr.find(qn('w:pgNumType'))
        fmt = pnt.get(qn('w:fmt')) if pnt is not None else 'None'
        start = pnt.get(qn('w:start')) if pnt is not None else 'None'
        print(f"  pgNumType: fmt={fmt}, start={start}")
        
        ftr = section.footer
        if ftr:
            print(f"  Footer paragraphs: {len(ftr.paragraphs)}")
            for j, p in enumerate(ftr.paragraphs):
                text = p.text.strip()
                instrs = [instr.text for instr in p._p.findall('.//' + qn('w:instrText'))]
                print(f"    Para {j}: text='{text}', instrText={instrs}")
        else:
            print("  No footer.")

check_footers(r'd:\RESEARCH PNGS\zukhruftheisis_cleaned.docx')

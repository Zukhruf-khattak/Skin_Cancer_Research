import docx
from docx.oxml.ns import qn

def inspect_v3():
    doc = docx.Document(r'd:\RESEARCH PNGS\zukhruftheisis_v3.docx')
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip().upper()
        if 'TABLE OF CONTENTS' in text or 'CONTENTS' in text:
            print(f"[{i}] {p.text.strip()}")
            # Print a few paragraphs before and after
            for j in range(max(0, i-5), min(len(doc.paragraphs), i+5)):
                print(f"  [{j}] {doc.paragraphs[j].text.strip()[:80]}")

inspect_v3()

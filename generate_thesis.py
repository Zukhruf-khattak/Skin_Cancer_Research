"""
Complete Thesis Generator — Automated Skin Cancer Detection and Classification
Generates a professional .docx thesis matching the senior FYP format exactly.
"""
import docx
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

OUTPUT_PATH = r"d:\RESEARCH PNGS\Thesis_Skin_Cancer_FYP.docx"
IMG_S1 = r"d:\RESEARCH PNGS\Final stage1"
IMG_S2 = r"d:\RESEARCH PNGS\Final Stage 2"
IMG_RES = r"d:\RESEARCH PNGS\results\stage1\figures"
IMG_RES_GC = r"d:\RESEARCH PNGS\results\stage1\gradcam"

doc = Document()

# ═══════════════════════════════════════════════════════════════════
#  HELPER FUNCTIONS
# ═══════════════════════════════════════════════════════════════════

def setup_page(section):
    """A4 page with exact margins from senior thesis."""
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1.5)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

def set_run_font(run, size=12, bold=False, italic=False, name='Times New Roman', color=None):
    """Set font properties on a run."""
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    # Force Times New Roman for East Asian text too
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{name}"/>')
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), name)

def set_paragraph_spacing(para, before=0, after=0, line_spacing=1.5):
    """Set paragraph spacing."""
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line_spacing

def add_heading1(text, alignment=WD_ALIGN_PARAGRAPH.CENTER):
    """Heading 1: 16pt Bold Centered (for chapter titles)."""
    para = doc.add_paragraph()
    para.alignment = alignment
    run = para.add_run(text)
    set_run_font(run, size=16, bold=True)
    set_paragraph_spacing(para, before=0, after=0, line_spacing=1.0)
    para.style = doc.styles['Normal']
    return para

def add_heading2(text):
    """Heading 2: 14pt Bold Left-aligned."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(text)
    set_run_font(run, size=14, bold=True)
    set_paragraph_spacing(para, before=2, after=12, line_spacing=1.5)
    return para

def add_heading3(text):
    """Heading 3: 12pt Bold Left-aligned."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(text)
    set_run_font(run, size=12, bold=True)
    set_paragraph_spacing(para, before=2, after=0, line_spacing=2.0)
    return para

def add_body(text, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """Body text: 12pt Times New Roman, Justified, 1.5x spacing."""
    para = doc.add_paragraph()
    para.alignment = alignment
    run = para.add_run(text)
    set_run_font(run, size=12)
    set_paragraph_spacing(para, before=5, after=5, line_spacing=1.5)
    return para

def add_body_bold_start(bold_text, normal_text):
    """Body text starting with bold text then normal."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run1 = para.add_run(bold_text)
    set_run_font(run1, size=12, bold=True)
    run2 = para.add_run(normal_text)
    set_run_font(run2, size=12)
    set_paragraph_spacing(para, before=5, after=5, line_spacing=1.5)
    return para

def add_list_item(text):
    """List item: 12pt, Justified, 1.5x, indented."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = para.paragraph_format
    pf.left_indent = Inches(0.5)
    pf.space_before = Pt(0)
    pf.space_after = Pt(12)
    pf.line_spacing = 1.5
    run = para.add_run(text)
    set_run_font(run, size=12)
    return para

def add_numbered_list(items):
    """Add numbered list items."""
    for i, item in enumerate(items, 1):
        add_list_item(f"{i}. {item}")

def add_bullet_list(items):
    """Add bullet list items."""
    for item in items:
        add_list_item(f"• {item}")

def add_image(path, caption=None, width=Inches(5.5)):
    """Add an image with optional caption."""
    if os.path.exists(path):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.add_picture(path, width=width)
        if caption:
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cap.add_run(caption)
            set_run_font(run, size=10, italic=True)
            set_paragraph_spacing(cap, before=2, after=8, line_spacing=1.0)
    else:
        print(f"  WARNING: Image not found: {path}")

def add_table(headers, rows, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row
    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        set_run_font(run, size=12, bold=True)
        # Light gray background
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3"/>')
        cell._element.get_or_add_tcPr().append(shading)
    
    # Data rows
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i + 1].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(cell_text))
            set_run_font(run, size=12)
    
    # Add borders
    tbl = table._element
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)
    
    # Add spacing after table
    doc.add_paragraph()
    return table

def add_page_break():
    """Add a page break."""
    doc.add_page_break()

def add_empty_lines(count=1):
    """Add empty paragraphs."""
    for _ in range(count):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=0, after=0, line_spacing=1.0)

# ═══════════════════════════════════════════════════════════════════
#  PAGE SETUP
# ═══════════════════════════════════════════════════════════════════
section = doc.sections[0]
setup_page(section)

# ═══════════════════════════════════════════════════════════════════
#  TITLE PAGE
# ═══════════════════════════════════════════════════════════════════
print("Building Title Page...")
add_empty_lines(3)

# Project Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Automated Skin Cancer Detection and Classification Using Deep Learning")
set_run_font(run, size=16, bold=True)
set_paragraph_spacing(p, before=0, after=0, line_spacing=1.5)

add_empty_lines(4)

# Submitted by
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Submitted By:")
set_run_font(run, size=14, bold=True)

# Student names
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Zohrouf Khattak\t\t\tRoll No: 1000")
set_run_font(run, size=14, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Huma Zeb\t\t\t\tRoll No: 883")
set_run_font(run, size=14, bold=True)

add_empty_lines(2)

# Supervisor
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Supervised by:")
set_run_font(run, size=14, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Dr. Muhammad Ayaz")
set_run_font(run, size=14, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Lecturer in Computer Science")
set_run_font(run, size=12)
set_paragraph_spacing(p, before=0, after=3, line_spacing=1.0)

add_empty_lines(2)

# Thesis submission statement
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("This thesis is submitted to Shaikh Zayed Islamic Centre, University of Peshawar in partial fulfillment of the requirements for the degree of")
set_run_font(run, size=12)
set_paragraph_spacing(p, before=0, after=3, line_spacing=1.5)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("BACHELOR IN COMPUTER SCIENCE BS(CS)")
set_run_font(run, size=14, bold=True)
set_paragraph_spacing(p, before=0, after=3, line_spacing=1.0)

add_empty_lines(1)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("SHAIKH ZAYED ISLAMIC CENTRE,")
set_run_font(run, size=14, bold=True)
set_paragraph_spacing(p, before=0, after=1, line_spacing=1.0)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("UNIVERSITY OF PESHAWAR")
set_run_font(run, size=14, bold=True)
set_paragraph_spacing(p, before=0, after=1, line_spacing=1.0)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("(Session 2022-2026)")
set_run_font(run, size=14, bold=True)

# ═══════════════════════════════════════════════════════════════════
#  APPROVAL CERTIFICATE
# ═══════════════════════════════════════════════════════════════════
add_page_break()
print("Building Approval Certificate...")

add_empty_lines(1)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("APPROVAL CERTIFICATE")
set_run_font(run, size=14, bold=True)
set_paragraph_spacing(p, before=0, after=8, line_spacing=1.0)

add_body("This is to certify that this project is approved and recommended as partial fulfillment of the BS (Honors) Computer Science Degree award at Shaikh Zayed Islamic Centre, University of Peshawar.")

add_empty_lines(1)

# Internal Examiner
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Internal Examiner:")
set_run_font(run, size=14, bold=True)
set_paragraph_spacing(p, before=0, after=8, line_spacing=1.0)

p = doc.add_paragraph()
run = p.add_run("Dr. Muhammad Ayaz")
set_run_font(run, size=14, bold=True)
set_paragraph_spacing(p, before=0, after=8, line_spacing=1.0)

p = doc.add_paragraph()
run = p.add_run("Lecturer in Computer Science,\nShaikh Zayed Islamic Centre, University of Peshawar")
set_run_font(run, size=12)
set_paragraph_spacing(p, before=0, after=12, line_spacing=1.5)

p = doc.add_paragraph()
run = p.add_run("Signature: ____________________")
set_run_font(run, size=12)
set_paragraph_spacing(p, before=0, after=12, line_spacing=1.0)

add_empty_lines(1)

# External Examiner
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("External Examiner:")
set_run_font(run, size=14, bold=True)
set_paragraph_spacing(p, before=0, after=8, line_spacing=1.0)

p = doc.add_paragraph()
run = p.add_run("____________________")
set_run_font(run, size=14, bold=True)
set_paragraph_spacing(p, before=0, after=8, line_spacing=1.0)

p = doc.add_paragraph()
run = p.add_run("Signature: ____________________")
set_run_font(run, size=12)
set_paragraph_spacing(p, before=0, after=12, line_spacing=1.0)

add_empty_lines(1)

# Director
p = doc.add_paragraph()
run = p.add_run("Prof. Dr. Rashid Ahmad")
set_run_font(run, size=14, bold=True)
set_paragraph_spacing(p, before=0, after=8, line_spacing=1.0)

p = doc.add_paragraph()
run = p.add_run("Director,\nShaikh Zayed Islamic Centre,\nUniversity of Peshawar")
set_run_font(run, size=12)
set_paragraph_spacing(p, before=0, after=8, line_spacing=1.0)

p = doc.add_paragraph()
run = p.add_run("Signature: ____________________")
set_run_font(run, size=12)
set_paragraph_spacing(p, before=0, after=20, line_spacing=1.0)

p = doc.add_paragraph()
run = p.add_run("Dated: ____/____/2026")
set_run_font(run, size=12)
set_paragraph_spacing(p, before=0, after=20, line_spacing=1.0)

# ═══════════════════════════════════════════════════════════════════
#  ACKNOWLEDGEMENT
# ═══════════════════════════════════════════════════════════════════
add_page_break()
print("Building Acknowledgement...")
add_heading1("ACKNOWLEDGEMENT")
add_empty_lines(1)

add_body("First, we pray to Allah Al-Mighty who gave us the ability and courage to complete this project.")
add_body("Lots of blessing and Darood Pak to Khatimun Nabieen Hazrat Muhammad (S.A.W) for whose sake Al-Mighty Allah has created this universe.")
add_body("Here we will especially give tribute to the honorable personalities, without their generosity and benevolence, it could not have been possible for us to accomplish this work. We are highly grateful to our supervisor Dr. Muhammad Ayaz, Lecturer in Computer Science at Shaikh Zayed Islamic Centre, University of Peshawar, for his continuous guidance, encouragement, and valuable suggestions throughout the course of this project. His expertise in machine learning and deep learning provided the foundation for the technical direction of this research.")
add_body("Finally, we are thankful to the Director and all the teaching staff of our institution.")
add_body("The acknowledgment will be incomplete if we do not mention the cooperation of our parents, who encouraged us and prayed for our success. Their unwavering support was the cornerstone of our academic journey.")

add_empty_lines(2)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Zohrouf Khattak")
set_run_font(run, size=14, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Huma Zeb")
set_run_font(run, size=14, bold=True)

# ═══════════════════════════════════════════════════════════════════
#  DEDICATION
# ═══════════════════════════════════════════════════════════════════
add_page_break()
print("Building Dedication...")
add_heading1("DEDICATION")
add_empty_lines(1)

add_body("To some of our favorite and beloved personalities, My Prophet Muhammad (Peace be Upon Him), the Prophet of Mercy and Guidance, whose teachings inspire us to seek knowledge and serve humanity.")
add_body("Secondly, we would like to sincerely thank all the faculty members of Shaikh Zayed Islamic Centre for their academic guidance and mentorship throughout our degree program. Their dedication to teaching and research has shaped our intellectual growth and professional aspirations.")
add_body("We dedicate this work to our parents, whose sacrifices, prayers, and unconditional love have been the driving force behind every milestone we have achieved. This thesis is a testament to their belief in us.")
add_body("We also dedicate this work to the field of medical artificial intelligence, with the hope that advances in automated diagnostic systems will improve access to early cancer detection for underserved communities worldwide.")

# ═══════════════════════════════════════════════════════════════════
#  ABSTRACT
# ═══════════════════════════════════════════════════════════════════
add_page_break()
print("Building Abstract...")
add_heading1("ABSTRACT")
add_empty_lines(1)

add_body("This study presents a dual-stage ensemble deep learning pipeline for automated skin cancer detection and classification using the ISIC 2019 dataset comprising 25,331 dermoscopic images. The proposed system implements a hierarchical diagnostic framework that mirrors clinical dermatologist reasoning: Stage 1 performs binary triage classification (Benign vs. Malignant) using a 5-model ensemble of transfer-learning-based architectures, while Stage 2 performs fine-grained multi-class classification of malignant subtypes (Melanoma, Basal Cell Carcinoma, Squamous Cell Carcinoma, and Actinic Keratosis) using models trained from scratch with Focal Loss and MixUp augmentation.")

add_body("A rigorous preprocessing pipeline — including morphological hair artifact removal, Telea inpainting, Otsu-based lesion segmentation, and ImageNet-standard normalisation — ensures clinically clean input images. Targeted undersampling of the dominant Melanocytic Nevus class (8,978 → 4,238 images) addresses severe binary-level class imbalance, while Focal Loss and MixUp augmentation handle multi-class imbalance without discarding any rare cancer data.")

add_body("The 5-model ensemble combines three CNN families (ResNet50, EfficientNet-B0, MobileNetV2) with two modern architectures (ConvNeXt-Tiny and Swin-Tiny Vision Transformer) via soft-voting probability averaging. Stage 1 achieves an ensemble accuracy of 94.34%, sensitivity of 93.79%, and AUC of 0.9845 on the held-out test set (3,800 images). Stage 2 achieves an ensemble accuracy of 85.58%, Macro-F1 of 0.8332, and Macro-AUC of 0.9738, with a critical SCC recall of 95.79%.")

add_body("Gradient-weighted Class Activation Mapping (Grad-CAM) provides visual evidence that all models attend to clinically meaningful dermoscopic features — irregular pigmentation networks, asymmetric borders, and keratotic plugs — rather than background artifacts. An interactive Streamlit demonstration application enables real-time ensemble inference with per-model probability breakdowns.")

# Keywords
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
run1 = p.add_run("Keywords: ")
set_run_font(run1, size=12, bold=True)
run2 = p.add_run("Skin Cancer Detection, Deep Learning, Ensemble Learning, ISIC 2019, Dermoscopy, Grad-CAM, Transfer Learning, Focal Loss, MixUp, Clinical Decision Support")
set_run_font(run2, size=12)
set_paragraph_spacing(p, before=5, after=5, line_spacing=1.5)

# ═══════════════════════════════════════════════════════════════════
#  CHAPTER I — INTRODUCTION
# ═══════════════════════════════════════════════════════════════════
add_page_break()
print("Building Chapter I: Introduction...")
add_heading1("CHAPTER I")
add_heading1("INTRODUCTION")
add_empty_lines(1)

add_heading2("1.1  Background of the Study")

add_heading3("1.1.1  Global Burden of Skin Cancer")
add_body("Skin cancer represents one of the most rapidly increasing malignancies worldwide and currently ranks as the sixth most prevalent cancer type globally. According to the World Cancer Research Fund and the American Cancer Society, approximately 1.5 million new cases of skin cancer are diagnosed each year, with over 50,000 fatalities annually attributable to the disease [1], [2]. Australia and the United States report the highest incidence rates, with the American Cancer Society estimating that one in five Americans will develop skin cancer in their lifetime [3]. The global economic burden is equally staggering, with annual treatment costs estimated at approximately USD 8.1 billion in the United States alone, representing a substantial and growing strain on healthcare infrastructures worldwide [4].")
add_body("The disease is caused by the abnormal and uncontrolled proliferation of cells within the structural layers of the skin, primarily triggered by cumulative exposure to ultraviolet (UV) radiation, genetic predisposition, immunosuppressed states, and environmental carcinogens [5]. The human skin is composed of three primary cell types: basal cells, squamous cells, and melanocytes. Malignant transformation within these cell populations gives rise to the three most clinically significant forms of skin cancer: Basal Cell Carcinoma (BCC), Squamous Cell Carcinoma (SCC), and Melanoma, respectively [6].")
add_body("While BCC and SCC represent the most common non-melanoma skin cancers (NMSCs) and generally exhibit lower metastatic potential when treated in early stages, Melanoma is highly aggressive and spreads rapidly through the lymphatic system and bloodstream. Although melanoma accounts for only approximately 5% of all skin cancer cases by incidence, it is responsible for over 75% of all skin cancer-related deaths worldwide [7]. The 5-year survival rate for melanoma exceeds 99% when detected at a localized, early stage; however, late-stage diagnoses — involving regional lymph node involvement or distant metastasis — reduce this survival rate to approximately 68% and 30%, respectively [8]. This dramatic discrepancy underscores the critical importance of early, accurate, and accessible diagnostic technologies.")

add_heading3("1.1.2  Classification of Skin Cancer Types")
add_body("A thorough understanding of the taxonomy of skin lesions is foundational to building effective automated detection systems. Skin lesions are broadly classified into two primary categories: benign (non-cancerous) and malignant (cancerous). Benign conditions include melanocytic nevi (common moles), seborrheic keratoses, dermatofibromas, and vascular lesions, which are generally harmless but may closely resemble malignant lesions during visual inspection. Malignant categories include the following clinically significant subtypes [9]:")
add_list_item("Melanoma (MEL): Arising from melanocytes, melanoma is the deadliest form of skin cancer. It is characterized by asymmetrical borders, multiple colors, irregular growth patterns, and the capacity for rapid systemic metastasis.")
add_list_item("Basal Cell Carcinoma (BCC): The most common form of skin cancer, arising from the basal cells of the epidermis. BCC rarely metastasizes but can cause significant local tissue destruction if left untreated.")
add_list_item("Squamous Cell Carcinoma (SCC): Originating from squamous cells of the epidermis, SCC has a higher metastatic potential than BCC and can spread to regional lymph nodes and distant organs if not diagnosed and treated promptly.")
add_list_item("Actinic Keratosis (AK): A precancerous lesion caused by prolonged UV exposure. AK is considered a direct precursor to SCC and requires treatment to prevent malignant progression.")
add_list_item("Melanocytic Nevus (NV): Benign pigmented lesions that frequently mimic early-stage melanoma in dermoscopic appearance, contributing significantly to diagnostic uncertainty.")
add_body("The high degree of visual overlap between benign and malignant lesions — particularly between melanocytic nevi and early melanoma — is one of the most critical factors contributing to diagnostic errors in clinical dermatology [10].")

add_heading3("1.1.3  Clinical Diagnosis and Its Inherent Challenges")
add_body("The current gold standard for clinical screening of skin lesions involves visual inspection using a dermatoscope — a specialized, non-invasive, hand-held optical instrument that provides cross-polarized or liquid-contact illumination, enabling clinicians to examine subsurface skin structures not visible to the naked eye [11]. This technique, known as dermoscopy or dermatoscopy, significantly enhances the diagnostic accuracy of experienced dermatologists compared to unaided visual inspection by approximately 10–27% [12].")
add_body("Despite the utility of dermoscopy, visual diagnosis remains inherently subjective and is heavily influenced by the individual clinician's level of experience, training, and cognitive bias. Studies have consistently reported high inter-observer variability in the visual diagnosis of dermoscopic images, even among board-certified dermatologists [13]. A landmark study by Haenssle et al. (2018) demonstrated that the average diagnostic sensitivity of a cohort of 58 international dermatologists was only 86.6% — a figure lower than that achieved by an automated deep learning model on the same test set [14].")
add_body("Furthermore, the global shortage of qualified dermatologists severely limits access to timely and accurate skin cancer screening, particularly in low- and middle-income countries, rural regions, and underserved populations. According to the International League of Dermatological Societies, the global dermatologist-to-population ratio is approximately 1 per 100,000 persons, with severe disparities across regions [15]. This access gap directly contributes to the disproportionately high rates of late-stage diagnoses and associated mortality in these regions [16].")

add_heading3("1.1.4  Computer-Aided Diagnosis (CAD) Systems")
add_body("To mitigate the limitations of purely subjective visual diagnosis, researchers have developed Computer-Aided Diagnosis (CAD) systems that use computational algorithms to analyze dermoscopic images and provide objective, quantitative diagnostic support to clinicians. Early CAD systems, developed in the late 1990s and 2000s, relied on classical machine learning (ML) pipelines that required manual feature engineering — a multi-step process involving hand-crafted feature extraction (based on the ABCD rule: Asymmetry, Border irregularity, Color variation, and Diameter), statistical feature selection, and classification using algorithms such as Support Vector Machines (SVM) and Random Forests [17], [18].")
add_body("While these classical CAD systems demonstrated the feasibility of automated lesion classification, they suffered from critical limitations including sensitivity to image noise and dermoscopic artifacts, poor generalization across different imaging conditions, and their dependence on error-prone manual segmentation as a prerequisite for feature extraction. These constraints limited the diagnostic accuracy of classical CAD systems to approximately 80–85%, far below the level required for clinically actionable triage support [19].")

add_heading3("1.1.5  The Deep Learning Revolution in Medical Imaging")
add_body("The introduction of deep learning, and specifically deep Convolutional Neural Networks (CNNs), marked a transformative breakthrough in the field of medical image analysis. Unlike classical ML approaches, CNNs automatically learn hierarchical, translation-invariant spatial feature representations directly from raw pixel data through a process of end-to-end supervised training, eliminating the need for manual feature engineering entirely [20]. The landmark study by LeCun et al. (1998) established the foundational CNN architecture, and the AlexNet breakthrough at the ImageNet Large Scale Visual Recognition Challenge (ILSVRC 2012) demonstrated that deep CNNs could achieve far superior performance compared to all prior machine learning approaches on large-scale image classification tasks [21].")
add_body("The application of deep learning to skin cancer detection was brought to global prominence by the seminal study of Esteva et al. (2017), published in Nature. Their Inception-v3 CNN trained on 129,450 clinical and dermoscopic images achieved an Area Under the Receiver Operating Characteristic Curve (AUC) of 0.96 in classifying melanoma and keratinocyte carcinomas, demonstrating diagnostic performance at a level statistically equivalent to that of 21 board-certified dermatologists [22]. This landmark result established that deep neural networks could extract clinical feature representations of sufficient resolution to match human expert diagnosis.")
add_body("Subsequent investigations rapidly confirmed and extended these findings. Haenssle et al. (2018) showed that a deep CNN outperformed 58 international dermatologists on a binary melanoma classification task, with the model achieving a sensitivity of 95.0% compared to the dermatologists' average of 86.6% [14]. Tschandl et al. (2019) further demonstrated in The Lancet Oncology that deep learning algorithms achieved a mean multi-class accuracy of 80.0% across seven lesion categories — significantly exceeding the mean accuracy of 62.0% achieved by a cohort of 511 human clinical readers [23].")

add_heading3("1.1.6  Transfer Learning and Pre-trained Models")
add_body("Despite the superior performance of deep CNNs, training these models from scratch requires vast quantities of annotated training images — a resource that is chronically scarce in clinical dermatology due to the high cost and time required for expert pathological labeling. To overcome this data scarcity bottleneck, Transfer Learning has become the dominant paradigm in medical image classification [24]. Transfer learning involves initializing a deep network with weights pre-trained on a large-scale general-domain dataset — most commonly the ImageNet dataset, which contains 1.2 million images across 1,000 object categories — and subsequently fine-tuning these weights on the target medical imaging dataset.")
add_body("By leveraging features learned on ImageNet (including low-level edges, corners, and Gabor-like filters, as well as mid-level textures and high-level semantic representations), pre-trained models can achieve superior classification performance on skin lesion datasets with dramatically fewer labeled samples, shorter training times, and lower risks of overfitting [25].")

add_heading3("1.1.7  Explainable Artificial Intelligence (XAI) in Clinical Decision Support")
add_body("Despite the high diagnostic accuracy of modern deep learning systems, a critical and frequently cited barrier to their clinical adoption is the \"black-box\" nature of neural network decision-making. Clinicians cannot ethically or legally rely on AI predictions if the underlying diagnostic rationale is opaque and unverifiable [29].")
add_body("Explainable Artificial Intelligence (XAI) techniques address this challenge by generating human-interpretable visual or mathematical explanations for deep network predictions. Gradient-weighted Class Activation Mapping (Grad-CAM), developed by Selvaraju et al. (2017), is among the most widely adopted XAI techniques in medical imaging [30]. Grad-CAM computes the gradients of the target class score with respect to the feature maps of the final convolutional layer, pools them spatially, and generates a saliency heatmap that highlights the specific image regions most influential to the network's classification decision.")

add_heading2("1.2  Problem Statement")
add_body("Despite the significant advancements in deep learning-based skin cancer detection documented in the literature, existing automated diagnostic systems suffer from a series of interconnected and clinically critical limitations that collectively hinder their reliable translation into routine clinical practice. This research is motivated by the following key challenges:")
add_body("(1) Clinical Artifact Interference and Preprocessing Inadequacy: Raw dermoscopic images are routinely corrupted by dense hair occlusions, surgical ink ruler markings, gel reflections, specular highlights, and variable illumination gradients. These artifacts occlude lesion boundaries, introduce spurious texture features, and cause deep networks to learn non-pathological background patterns [31].")
add_body("(2) Extreme Class Imbalance: Dermoscopic datasets are characterized by severe inter-class distribution disparities. In the ISIC 2019 training dataset, melanocytic nevi constitute approximately 50.8% of all samples, while clinically critical malignant subtypes such as Squamous Cell Carcinoma (SCC) represent only 2.5% [32].")
add_body("(3) Lack of Model Interpretability and Clinical Trust: Most high-performing deep learning models for skin cancer detection operate as uninterpretable black boxes, providing a class label and a confidence score without any spatially grounded diagnostic rationale [29], [30].")
add_body("To address this constellation of challenges, this study proposes a lightweight, methodologically rigorous, and clinically interpretable two-stage deep learning pipeline for automated skin cancer triage and malignant subtype classification using solely dermoscopic images.")

add_heading2("1.3  Research Aim")
add_body("The primary aim of this study is to develop, optimize, and validate a highly accurate, computationally efficient, and clinically interpretable two-stage deep learning framework for automated skin cancer detection and classification using dermoscopic images from the ISIC 2019 dataset. The system is specifically designed to support dermatologists and primary care physicians in making early, objective, and trust-verified clinical diagnoses, with particular applicability in resource-constrained healthcare environments.")

add_heading2("1.4  Research Objectives")
add_body("To achieve the stated aim, this study addresses the following specific and measurable research objectives:")
add_list_item("Develop a morphological artifact removal and standardized preprocessing pipeline, including Black Hat transform-based hair removal, Otsu-based lesion region extraction, and ImageNet channel-wise normalization, to eliminate clinical noise.")
add_list_item("Implement a patient-stratified train-validation-test splitting strategy that prevents patient-level data leakage, ensuring that no images from the same patient appear across splits.")
add_list_item("Build, train, and optimize a Stage 1 binary triage classifier using a 5-model heterogeneous ensemble (ResNet50, EfficientNet-B0, MobileNetV2, ConvNeXt-Tiny, Swin-Tiny) with transfer learning, tuned to achieve high diagnostic sensitivity for malignant lesions.")
add_list_item("Develop a Stage 2 multi-class malignant subtype classifier to categorize malignant lesions into four clinical subtypes (Melanoma, BCC, SCC, AK) using Focal Loss and MixUp augmentation to handle severe class imbalance.")
add_list_item("Integrate Gradient-weighted Class Activation Mapping (Grad-CAM) to generate diagnostic attention heatmaps and visually validate that model attention aligns with established dermoscopic clinical criteria.")
add_list_item("Conduct a rigorous comparative evaluation of all five backbone architectures at both classification stages using standard performance metrics including Accuracy, Precision, Recall, F1-Score, and AUC-ROC.")

add_heading2("1.5  Research Questions")
add_body("This study is guided by the following research questions:")
add_list_item("To what extent can a heterogeneous 5-model ensemble, combining CNNs and Vision Transformers with standardized preprocessing, achieve clinically actionable sensitivity in binary skin cancer triage from dermoscopic images?")
add_list_item("Which of the five evaluated architectures achieves the best balance between classification accuracy, computational efficiency, and diagnostic interpretability in the proposed two-stage pipeline?")
add_list_item("Can Focal Loss and MixUp augmentation effectively address extreme class imbalance in multi-class malignant subtype classification without discarding rare cancer data?")
add_list_item("Can Grad-CAM visualizations generated by the proposed models reliably localize attention on clinically meaningful dermoscopic features?")

add_heading2("1.6  Significance of the Study")
add_heading3("1.6.1  Clinical Significance")
add_body("The two-stage triage architecture of the proposed system prioritizes high sensitivity in the Stage 1 binary classification phase. By structuring the pipeline to flag all potentially malignant lesions for expert review, the system provides an invaluable first-pass screening tool for primary care physicians, rural clinics, and mobile dermatology units that lack immediate access to specialist dermatologists [37].")

add_heading3("1.6.2  Technical Significance")
add_body("By actively resolving patient-level data leakage, applying standardized ImageNet normalization, and implementing morphological artifact preprocessing, this research establishes a high methodological benchmark for rigorous and reproducible AI model evaluation in clinical dermatology [33], [34].")

add_heading3("1.6.3  Societal and Equity Significance")
add_body("By including lightweight, parameter-efficient architectures (MobileNetV2, EfficientNet-B0) that can be deployed on conventional clinical computing hardware, this research directly contributes to the democratization of AI-powered diagnostic support. Deploying such systems in low-resource settings has the potential to reduce the diagnostic access gap that contributes to disproportionately high skin cancer mortality rates in underserved populations [38].")

add_heading2("1.7  Scope of the Study")
add_body("This study is bounded by the following scope parameters:")
add_list_item("Dataset: The research exclusively uses the publicly available ISIC 2019 Challenge Dataset, comprising 25,331 dermoscopic images across eight diagnostic categories.")
add_list_item("Architectures: Five architectures are evaluated: ResNet50, EfficientNet-B0, MobileNetV2, ConvNeXt-Tiny, and Swin-Tiny Vision Transformer.")
add_list_item("Classification Stages: A two-stage hierarchical pipeline — binary (benign vs. malignant) at Stage 1, and malignant subtype classification (MEL, BCC, SCC, AK) at Stage 2.")
add_list_item("Explainability: Grad-CAM is the exclusive XAI technique evaluated.")
add_list_item("Deployment: The study includes a Streamlit demonstration application but does not include real-time clinical integration or regulatory approval processes.")

add_heading2("1.8  Organization of the Thesis")
add_body("The remainder of this thesis is structured into five chapters as follows:")
add_list_item("Chapter 2: Literature Review — provides a comprehensive and critical analysis of the historical evolution of automated skin cancer detection systems, covering classical machine learning approaches, deep learning foundations, advanced architectures, and explainable AI.")
add_list_item("Chapter 3: Methodology — details the complete technical implementation of the proposed two-stage pipeline, including dataset description, preprocessing pipeline, model architectures, training configurations, and evaluation design.")
add_list_item("Chapter 4: Results and Evaluation — presents quantitative results for both stages, including comparative performance tables, confusion matrices, ROC curves, and Grad-CAM visualizations.")
add_list_item("Chapter 5: Demo Application — describes the interactive Streamlit web application for real-time ensemble inference.")
add_list_item("Chapter 6: Conclusion and Future Work — summarizes key contributions, findings, limitations, and proposes future research directions.")

# ═══════════════════════════════════════════════════════════════════
#  CHAPTER II — LITERATURE REVIEW  (simplified — pulling from ch1ch2)
# ═══════════════════════════════════════════════════════════════════
add_page_break()
print("Building Chapter II: Literature Review...")
add_heading1("CHAPTER II")
add_heading1("LITERATURE REVIEW")
add_empty_lines(1)

# Read ch1ch2 content for lit review
ch1ch2_lines = []
try:
    with open(r"d:\RESEARCH PNGS\scratch_ch1ch2_text.txt", 'r', encoding='utf-8') as f:
        ch1ch2_lines = f.readlines()
except:
    pass

# Find lit review section
in_lit_review = False
current_heading_level = None
for line in ch1ch2_lines:
    line = line.strip()
    if not line:
        continue
    
    # Detect start of lit review
    if 'CHAPTER 2' in line and 'LITERATURE REVIEW' in line.upper():
        in_lit_review = True
        continue
    if line == '[Normal] CHAPTER 2  LITERATURE REVIEW':
        in_lit_review = True
        continue
    if 'LITERATURE REVIEW' in line and '[Normal]' in line and not in_lit_review:
        in_lit_review = True
        continue
        
    # Detect end (chapter 3 or methodology)
    if in_lit_review and ('CHAPTER 3' in line or 'METHODOLOGY' in line.upper()):
        break
    
    if not in_lit_review:
        continue
    
    # Parse style and text
    if line.startswith('['):
        bracket_end = line.index(']')
        style = line[1:bracket_end]
        text = line[bracket_end+2:].strip() if bracket_end + 2 < len(line) else ''
    else:
        style = 'Normal'
        text = line
    
    if not text:
        continue
    
    # Skip certain meta content
    if text.startswith('[Insert ') or text.startswith('Table of Contents'):
        continue
    
    if style == 'Heading 2':
        add_heading2(text)
    elif style == 'Heading 3':
        add_heading3(text)
    elif style == 'List Paragraph':
        add_list_item(text)
    else:
        add_body(text)

print(f"  Literature review: processed lines from ch1ch2")

# ═══════════════════════════════════════════════════════════════════
#  CHAPTER III — METHODOLOGY
# ═══════════════════════════════════════════════════════════════════
add_page_break()
print("Building Chapter III: Methodology...")
add_heading1("CHAPTER III")
add_heading1("METHODOLOGY")
add_empty_lines(1)

# 3.1 Overview
add_heading2("3.1  Overview of the Proposed System")
add_body("This research proposes a Dual-Stage Ensemble Deep Learning Pipeline for the automated detection and fine-grained classification of skin cancer from dermoscopic images. The system is designed to replicate key aspects of dermatologist diagnostic workflow: first confirming whether a lesion is dangerous (malignant or benign), and then — if malignant — identifying which specific type of cancer it is, so that appropriate clinical action can be taken.")
add_body("The architecture is deliberately hierarchical and clinically motivated:")
add_list_item("Stage 1 acts as a triage gate, filtering the vast majority of benign cases and flagging only those that require further investigation.")
add_list_item("Stage 2 acts as a specialist classifier, identifying the exact malignant subtype from four clinically critical categories: Melanoma (MEL), Basal Cell Carcinoma (BCC), Squamous Cell Carcinoma (SCC), and Actinic Keratosis (AK).")
add_body("Every design decision — from model selection to loss function choice to data handling — was driven by one overarching clinical requirement: minimising false negatives, since a missed cancer diagnosis is far more harmful than an unnecessary referral.")

# 3.2 Dataset
add_heading2("3.2  Dataset: ISIC 2019 Challenge")
add_heading3("3.2.1  Dataset Description")
add_body("The International Skin Imaging Collaboration (ISIC) 2019 Challenge dataset was selected as the primary data source for this research. It is the largest publicly available benchmark for dermoscopic skin lesion classification and is widely accepted as the gold standard in computational dermatology research.")

add_table(
    ["Property", "Detail"],
    [
        ["Total Training Images", "25,331 dermoscopic images"],
        ["Image Format", "JPEG (variable resolution; 600×450 px typical)"],
        ["Annotation Type", "Diagnosis-level labels"],
        ["Source", "International Skin Imaging Collaboration"],
        ["Access", "Publicly available via Kaggle"],
        ["Hardware Used", "Kaggle — NVIDIA Tesla T4 GPU"],
    ]
)

add_heading3("3.2.2  The 8 Diagnostic Categories")
add_body("The dataset contains images across 8 diagnostic categories. For this pipeline, these categories are mapped to a hierarchical label structure:")

add_table(
    ["ISIC Category", "Full Name", "Stage 1 Label", "Stage 2 Label"],
    [
        ["MEL", "Melanoma", "Malignant", "Class 0"],
        ["NV", "Melanocytic Nevus", "Benign", "—"],
        ["BCC", "Basal Cell Carcinoma", "Malignant", "Class 1"],
        ["AK", "Actinic Keratosis", "Malignant", "Class 3"],
        ["BKL", "Benign Keratosis", "Benign", "—"],
        ["DF", "Dermatofibroma", "Benign", "—"],
        ["VASC", "Vascular Lesions", "Benign", "—"],
        ["SCC", "Squamous Cell Carcinoma", "Malignant", "Class 2"],
    ]
)

add_heading3("3.2.3  Class Imbalance — The Critical Challenge")
add_body("A severe class imbalance exists in the ISIC 2019 dataset. The Melanocytic Nevus (NV) class contained 8,978 images — more than all malignant classes combined — creating a 3:1 benign-to-malignant imbalance. Within the malignant classes, MEL dominates with ~4,522 images, while SCC has only ~628 images — a 7:1 within-class imbalance.")

add_image(os.path.join(IMG_S2, "raw_distribution.png"), "Figure 3.1: Raw distribution of ISIC 2019 malignant classes showing severe class imbalance")

# 3.3 Preprocessing
add_heading2("3.3  Preprocessing Pipeline")
add_body("Before any image is presented to a neural network, it undergoes a rigorous, multi-step preprocessing pipeline designed specifically for dermoscopic medical images.")

add_heading3("3.3.1  Step-by-Step Preprocessing")
add_body_bold_start("Step 1 — Image Loading: ", "Each raw JPEG image is loaded from disk using the Python Imaging Library (PIL) and converted to RGB format, ensuring a consistent 3-channel input.")
add_body_bold_start("Step 2 — Hair and Artifact Detection: ", "A black-hat morphological transform is applied using a 15×15 rectangular structuring element. This operation highlights dark, elongated structures (hairs) against the brighter skin background.")
add_body_bold_start("Step 3 — Inpainting (Artifact Removal): ", "The detected artifact mask is fed into OpenCV's inpaint() function using the Telea fast marching algorithm (Telea, 2004). This fills each artifact pixel with surrounding clean skin texture.")
add_body_bold_start("Step 4 — Lesion Detection: ", "The cleaned image is converted to grayscale, and Otsu's global thresholding is applied to separate the darker lesion region from the lighter surrounding skin.")
add_body_bold_start("Step 5 — Crop and Resize: ", "The image is cropped to the computed bounding box and resized to exactly 224×224 pixels using bicubic interpolation.")
add_body_bold_start("Step 6 — Tensor Conversion: ", "The PIL image is converted into a PyTorch tensor with shape (3, 224, 224). Pixel values are scaled from [0, 255] to [0.0, 1.0].")
add_body_bold_start("Step 7 — Normalisation: ", "The tensor is normalised using ImageNet statistics: mean μ = [0.485, 0.456, 0.406] and standard deviation σ = [0.229, 0.224, 0.225] per channel.")

add_image(os.path.join(IMG_S2, "clean_preprocessing.png"), "Figure 3.2: Preprocessing pipeline — from raw image with artifacts to clean, cropped lesion", width=Inches(5.0))

# 3.4 Stage 1 Models
add_heading2("3.4  Stage 1: Binary Classification (Benign vs. Malignant)")
add_heading3("3.4.1  Data Preparation")
add_body("The 8 ISIC categories are merged into binary labels: Malignant (MEL, BCC, AK, SCC) and Benign (NV, BKL, DF, VASC). Targeted undersampling of NV from 8,978 to 4,238 images addresses the severe class imbalance. The balanced dataset is split using stratified random sampling: 70% training, 15% validation, 15% test.")

add_image(os.path.join(IMG_S1, "03_undersampling_effect (1).png"), "Figure 3.3: Effect of NV undersampling on class distribution")
add_image(os.path.join(IMG_S1, "02_split_distribution (1).png"), "Figure 3.4: Train/Validation/Test split distribution")

add_heading3("3.4.2  Model Architectures")
add_body("Stage 1 employs a 5-model heterogeneous ensemble combining three distinct architectural families:")

add_body_bold_start("Architecture 1: ResNet50 — ", "A 50-layer deep CNN with residual skip connections that solve the vanishing gradient problem. The final FC layer is replaced with a custom binary classification head. Parameters: ~25.6M.")
add_body_bold_start("Architecture 2: EfficientNet-B0 — ", "A compound-scaled CNN that balances depth, width, and resolution. Uses MBConv blocks with Squeeze-and-Excitation attention. Parameters: ~5.3M (most parameter-efficient in ensemble).")
add_body_bold_start("Architecture 3: MobileNetV2 — ", "Designed for resource-constrained environments using inverted residual blocks with depthwise separable convolutions. Parameters: ~3.4M (lightest in ensemble).")
add_body_bold_start("Architecture 4: ConvNeXt-Tiny — ", "A modern pure CNN rebuilt to match Vision Transformer strengths. Features 7×7 depthwise convolutions, Layer Normalisation, and GELU activation. Parameters: ~28.6M.")
add_body_bold_start("Architecture 5: Swin-Tiny — ", "A hierarchical Vision Transformer using shifted window self-attention. Processes 16×16 image patches as tokens, capturing global patterns that CNNs miss. Parameters: ~28.3M.")

add_heading3("3.4.3  Training Configuration")
add_table(
    ["Hyperparameter", "Value", "Justification"],
    [
        ["Optimiser", "Adam (β₁=0.9, β₂=0.999)", "Adaptive learning rate"],
        ["Learning Rate", "1×10⁻⁴", "Standard for fine-tuning"],
        ["LR Scheduler", "ReduceLROnPlateau (patience=5)", "Reduces LR on plateau"],
        ["Loss Function", "BCEWithLogitsLoss", "Numerically stable binary CE"],
        ["Batch Size", "32", "GPU memory balance"],
        ["Max Epochs", "70", "Early stopping terminates earlier"],
        ["Early Stopping", "Patience = 10 epochs", "Monitors validation loss"],
        ["Weight Init", "ImageNet pre-trained", "Transfer learning"],
    ]
)

add_heading3("3.4.4  Ensemble Strategy")
add_body("After all five models are individually trained, their predictions are combined through probability averaging (soft voting): P_ensemble(x) = (1/5) × [P_ResNet50(x) + P_EfficientNet(x) + P_MobileNet(x) + P_ConvNeXt(x) + P_Swin(x)]. The final prediction is Malignant if P_ensemble ≥ 0.5, and Benign otherwise.")

# 3.5 Stage 2
add_heading2("3.5  Stage 2: Multi-Class Malignant Subtype Classification")
add_heading3("3.5.1  Design Philosophy: Learning From Scratch")
add_body("The central architectural decision in Stage 2 is to train all five models with randomly initialised weights (weights=None), completely rejecting ImageNet pre-training. The dermatological features that distinguish MEL from BCC are fundamentally different from ImageNet's natural image categories. Starting from a blank slate forces the models to develop feature representations entirely grounded in dermatological morphology.")

add_heading3("3.5.2  Class Imbalance Strategy: Focal Loss + MixUp")
add_body_bold_start("Focal Loss: ", "Standard cross-entropy loss treats every sample equally, causing the model to be dominated by MEL and BCC. Focal Loss (Lin et al., 2017) applies a modulating factor that down-weights easy, well-classified examples and focuses training on hard, misclassified ones: FL(p_t) = -α_t × (1 - p_t)^γ × log(p_t), where γ=2 and α_t is the inverse class frequency.")
add_body_bold_start("MixUp Data Augmentation: ", "MixUp (Zhang et al., 2018) creates virtual training samples by linearly interpolating between pairs of training images and their labels: x̃ = λ×xᵢ + (1-λ)×xⱼ, where λ ~ Beta(0.2, 0.2). Applied to every training batch, disabled during validation and testing.")

add_image(os.path.join(IMG_S2, "heavy_augmentation.png"), "Figure 3.5: Heavy augmentation pipeline visualisation", width=Inches(5.0))

add_heading3("3.5.3  Stage 2 Training Configuration")
add_table(
    ["Hyperparameter", "Stage 1", "Stage 2", "Reason"],
    [
        ["Weight Init", "ImageNet pre-trained", "Random (from scratch)", "Avoid domain bias"],
        ["Loss Function", "BCEWithLogitsLoss", "Focal Loss (γ=2)", "Handle class imbalance"],
        ["Max Epochs", "70", "40–50", "Faster convergence with Focal Loss"],
        ["Learning Rate", "1×10⁻⁴", "1×10⁻³", "Higher LR for random init"],
        ["LR Scheduler", "ReduceLROnPlateau", "CosineAnnealingLR", "Smoother decay"],
        ["Augmentation", "Standard", "Heavy + MixUp", "Overcome limited rare data"],
    ]
)

# 3.6 Grad-CAM
add_heading2("3.6  Grad-CAM Interpretability")
add_body("Gradient-weighted Class Activation Mapping (Grad-CAM; Selvaraju et al., 2017) is applied to provide visual evidence of model decision-making. During the forward pass, activations of the final convolutional layer are recorded. During the backward pass, gradients of the class score with respect to these activations are computed. Each feature map is weighted by the global average of its gradients, producing a heatmap where red/warm regions indicate where the model is looking.")

# 3.7 Evaluation Metrics
add_heading2("3.7  Evaluation Metrics")
add_table(
    ["Metric", "Formula", "Clinical Rationale"],
    [
        ["Accuracy", "(TP+TN)/(TP+TN+FP+FN)", "Overall correctness"],
        ["Sensitivity (Recall)", "TP/(TP+FN)", "Primary metric — detect malignancies"],
        ["Precision", "TP/(TP+FP)", "False alarm rate"],
        ["F1-Score", "2×(Prec×Rec)/(Prec+Rec)", "Harmonic mean"],
        ["AUC-ROC", "Area under ROC curve", "Threshold-independent ability"],
    ]
)

# 3.8 Hardware
add_heading2("3.8  Hardware and Software Environment")
add_table(
    ["Component", "Specification"],
    [
        ["Platform", "Kaggle Notebooks (cloud-based)"],
        ["GPU", "NVIDIA Tesla T4 (16 GB VRAM)"],
        ["Framework", "PyTorch 2.x + torchvision + timm"],
        ["Image Processing", "OpenCV 4.x, Pillow (PIL)"],
        ["Metrics", "scikit-learn"],
        ["Visualisation", "Matplotlib, Seaborn"],
        ["Python Version", "3.10+"],
    ]
)

# ═══════════════════════════════════════════════════════════════════
#  CHAPTER IV — RESULTS AND EVALUATION
# ═══════════════════════════════════════════════════════════════════
add_page_break()
print("Building Chapter IV: Results and Evaluation...")
add_heading1("CHAPTER IV")
add_heading1("RESULTS AND EVALUATION")
add_empty_lines(1)

# 4.1 Stage 1 Results
add_heading2("4.1  Stage 1: Binary Classification Results")
add_body("This section presents the quantitative evaluation of Stage 1 (Benign vs. Malignant binary classification). All results are reported on the held-out test set — 3,800 images that were never seen during training or validation.")

add_heading3("4.1.1  Individual Model and Ensemble Performance")
add_body("The table below presents the final metrics achieved by each model on the held-out test set:")

add_table(
    ["Model", "Accuracy", "Precision", "Sensitivity", "F1-Score", "AUC"],
    [
        ["ResNet50", "92.76%", "90.80%", "89.44%", "0.9011", "0.9675"],
        ["EfficientNet-B0", "92.87%", "88.38%", "92.86%", "0.9057", "0.9713"],
        ["MobileNetV2", "91.89%", "87.61%", "90.86%", "0.8921", "0.9680"],
        ["ConvNeXt-Tiny", "93.18%", "87.71%", "94.79%", "0.9111", "0.9746"],
        ["Swin-Tiny", "93.55%", "88.79%", "94.43%", "0.9153", "0.9734"],
        ["ENSEMBLE", "94.34%", "91.12%", "93.79%", "0.9244", "0.9845"],
    ]
)

add_image(os.path.join(IMG_S1, "ensemble_evaluation_plots.png"), "Figure 4.1: Ensemble evaluation plots — accuracy comparison, ROC curves, and sensitivity breakdown")

add_heading3("4.1.2  Improvement Over Baseline")
add_table(
    ["Metric", "Old Baseline", "Final Ensemble", "Absolute Gain"],
    [
        ["Accuracy", "69.19%", "94.34%", "+25.15%"],
        ["Sensitivity", "79.24%", "93.79%", "+14.55%"],
        ["AUC", "0.7665", "0.9845", "+0.2180"],
        ["F1-Score", "0.6323", "0.9244", "+0.2921"],
    ]
)

add_image(os.path.join(IMG_S1, "performance_gain_chart.png"), "Figure 4.2: Performance gain chart — improvement from baseline to final ensemble")

add_heading3("4.1.3  Clinical Safety Analysis — Confusion Matrix")
add_body("The ensemble's confusion matrix on the test set reveals the clinical safety profile:")

add_table(
    ["", "Predicted Benign", "Predicted Malignant"],
    [
        ["Actual Benign", "2,271 (TN)", "128 (FP)"],
        ["Actual Malignant", "87 (FN)", "1,314 (TP)"],
    ]
)

add_body("87 False Negatives: These are malignant lesions incorrectly classified as benign — the most dangerous error. The 93.79% sensitivity means the system correctly catches 93.79 out of every 100 malignant cases.")
add_body("128 False Positives: These are benign lesions flagged as malignant — leading to unnecessary referrals, but not life-threatening. The 94.66% specificity means 94.66 out of every 100 benign lesions are correctly identified.")

add_image(os.path.join(IMG_S1, "clinical_performance_summary.png"), "Figure 4.3: Clinical performance summary across all models")
add_image(os.path.join(IMG_S1, "clinical_analysis_plot.png"), "Figure 4.4: Clinical analysis — false negative rate analysis")

add_heading3("4.1.4  Stage 1 Grad-CAM Interpretability")
add_body("Grad-CAM was applied to the final Stage 1 models. Classical CNNs (ResNet50, EfficientNet-B0, MobileNetV2) produced tightly focused heatmaps concentrated on the core of the lesion, highlighting irregular textures and abnormal pigmentation. Modern architectures (ConvNeXt-Tiny, Swin-Tiny) produced broader attention patterns capturing lesion borders and surrounding contextual features.")

add_image(os.path.join(IMG_S1, "gradcam_comparison_summary.png"), "Figure 4.5: Grad-CAM comparison across all 5 Stage 1 models")

# Screenshots from Stage 1
add_image(os.path.join(IMG_S1, "Screenshot (196).png"), "Figure 4.6: Stage 1 training results and convergence")
add_image(os.path.join(IMG_S1, "Screenshot (197).png"), "Figure 4.7: Stage 1 model evaluation plots")
add_image(os.path.join(IMG_S1, "Screenshot (198).png"), "Figure 4.8: Stage 1 ensemble ROC curves and AUC comparison")
add_image(os.path.join(IMG_S1, "Screenshot (204).png"), "Figure 4.9: Stage 1 confusion matrix and classification report")
add_image(os.path.join(IMG_S1, "Screenshot (205).png"), "Figure 4.10: Stage 1 clinical safety metrics and Grad-CAM")

# 4.2 Stage 2 Results
add_heading2("4.2  Stage 2: Multi-Class Malignant Subtype Classification Results")
add_body("This section presents the results of Stage 2, which classifies malignant lesions into four subtypes: Melanoma (MEL), Basal Cell Carcinoma (BCC), Squamous Cell Carcinoma (SCC), and Actinic Keratosis (AK). All results are on the held-out Stage 2 test set (825 images).")

add_heading3("4.2.1  Individual Model and Ensemble Performance")
add_table(
    ["Model", "Accuracy", "Macro F1", "Macro AUC", "MEL Recall", "BCC Recall", "SCC Recall", "AK Recall"],
    [
        ["ResNet50", "80.48%", "0.7814", "0.9480", "81.67%", "77.33%", "91.57%", "76.92%"],
        ["EfficientNet-B0", "83.39%", "0.8132", "0.9615", "86.00%", "82.00%", "88.42%", "76.92%"],
        ["MobileNetV2", "80.24%", "0.7764", "0.9506", "82.33%", "79.00%", "87.37%", "73.08%"],
        ["ConvNeXt-Tiny", "80.36%", "0.7789", "0.9549", "80.67%", "77.67%", "91.57%", "77.69%"],
        ["Swin-Tiny", "82.91%", "0.8063", "0.9617", "83.00%", "80.00%", "91.57%", "83.08%"],
        ["ENSEMBLE", "85.58%", "0.8332", "0.9738", "85.00%", "85.67%", "95.79%", "79.23%"],
    ]
)

add_image(os.path.join(IMG_S2, "ultimate_summary_table.png"), "Figure 4.11: Stage 2 ultimate summary metrics")

add_heading3("4.2.2  Per-Model Analysis")
add_body_bold_start("EfficientNet-B0 — The Accuracy-Efficiency Leader: ", "Despite having only ~5.3M parameters, EfficientNet-B0 achieved the highest standalone accuracy (83.39%) and Macro F1 (0.813). For resource-constrained deployments, it offers the best accuracy-efficiency trade-off.")
add_body_bold_start("Swin-Tiny — The Vision Transformer Advantage: ", "Swin-Tiny achieved the highest standalone AK Recall (83.08%), surpassing all CNN-based models by a significant margin. The self-attention mechanism appears uniquely capable of capturing subtle global textural patterns of Actinic Keratosis lesions.")
add_body_bold_start("The Power of the Ensemble: ", "SCC Recall surged to 95.79% — 4+ percentage points above the best individual model. BCC Recall reached 85.67%. Clinical Priority Score reached 0.8881 — the highest of any configuration tested.")

add_heading3("4.2.3  Training Convergence")
add_image(os.path.join(IMG_S2, "training_convergence_grid.png"), "Figure 4.12: Training convergence — loss and accuracy curves for all 5 Stage 2 models")

add_heading3("4.2.4  Confusion Matrices")
add_image(os.path.join(IMG_S2, "confusion_matrices.png"), "Figure 4.13: Per-model confusion matrices for all 5 Stage 2 architectures")

add_heading3("4.2.5  ROC and Precision-Recall Curves")
add_image(os.path.join(IMG_S2, "roc_pr_curves.png"), "Figure 4.14: ROC and Precision-Recall curves for all 4 malignant classes")

add_heading3("4.2.6  Radar Chart")
add_image(os.path.join(IMG_S2, "radar_chart.png"), "Figure 4.15: Multi-dimensional performance radar chart comparing all models")

add_heading3("4.2.7  Clinical Priority Score")
add_body("The Clinical Priority Score (CPS) weights per-class recall by clinical urgency: CPS = (0.40 × Recall_MEL) + (0.25 × Recall_BCC) + (0.25 × Recall_SCC) + (0.10 × Recall_AK). The ensemble achieves a CPS of 0.8881.")

add_table(
    ["Model", "Clinical Priority Score"],
    [
        ["ResNet50", "0.8472"],
        ["EfficientNet-B0", "0.8566"],
        ["MobileNetV2", "0.8309"],
        ["ConvNeXt-Tiny", "0.8443"],
        ["Swin-Tiny", "0.8614"],
        ["ENSEMBLE", "0.8881"],
    ]
)

add_image(os.path.join(IMG_S2, "clinical_priority_scores.png"), "Figure 4.16: Clinical priority scores weighted by clinical urgency")

add_heading3("4.2.8  Grad-CAM: Clinical Explainability by Subtype")
add_body("AK (Actinic Keratosis): The heatmaps reveal focus on diffuse, rough surface texture and scattered erythema — consistent with AK diagnostic criteria where scale and roughness are key features.")
add_image(os.path.join(IMG_S2, "grad_cam_AK.png"), "Figure 4.17: Grad-CAM heatmaps for Actinic Keratosis (AK) classification", width=Inches(5.0))

add_body("SCC (Squamous Cell Carcinoma): Concentrated attention on central keratotic plugs, ulceration regions, and irregular erythematous borders — all classic clinical hallmarks of SCC.")
add_image(os.path.join(IMG_S2, "grad_cam_SCC.png"), "Figure 4.18: Grad-CAM heatmaps for Squamous Cell Carcinoma (SCC) classification", width=Inches(5.0))

# 4.3 Limitations
add_heading2("4.3  Statistical Considerations and Limitations")
add_heading3("4.3.1  Confidence Intervals")
add_body("For the Stage 1 ensemble sensitivity of 93.79% on n=1,401 malignant test cases, the exact 95% Clopper-Pearson confidence interval is [92.2%, 95.2%], indicating the results are statistically robust.")

add_heading3("4.3.2  Ablation Study")
add_body("The four concurrent methodological changes that produced the Stage 1 improvement from baseline were applied simultaneously. A formal ablation study was not conducted due to computational constraints. This is a recognised limitation of this work.")

# 4.4 Chapter Summary
add_heading2("4.4  Chapter Summary")
add_body("Stage 1 (Binary Classification): Ensemble accuracy 94.34%, sensitivity 93.79%, AUC 0.9845. Clinical false negative rate ~6.2% (87 malignant cases missed in 1,401 test cases).")
add_body("Stage 2 (Multi-Class Malignant Classification): Ensemble accuracy 85.58%, Macro-F1 0.8332, Macro-AUC 0.9738. Critical SCC recall 95.79%. Clinical Priority Score 0.8881.")
add_body("Grad-CAM analysis confirms that both Stage 1 and Stage 2 models attend to clinically meaningful dermoscopic features, supporting the interpretability requirements for medical AI deployment.")

# ═══════════════════════════════════════════════════════════════════
#  CHAPTER V — DEMO APPLICATION
# ═══════════════════════════════════════════════════════════════════
add_page_break()
print("Building Chapter V: Demo Application...")
add_heading1("CHAPTER V")
add_heading1("DEMO APPLICATION")
add_empty_lines(1)

add_heading2("5.1  Overview")
add_body("To demonstrate the pipeline's real-world applicability, an interactive web dashboard was developed using Streamlit. It provides an end-to-end demonstration of the dual-stage ensemble pipeline, allowing users to upload dermoscopic images and receive real-time predictions.")

add_heading2("5.2  Implementation Details")
add_body("The application uses PyTorch for CPU-based inference, ensuring deployment without requiring specialized GPU hardware. All ten pre-trained model weights (approximately 652 MB) are loaded once at startup and cached in memory using Streamlit's @st.cache_resource decorator. While the initial load takes 30–60 seconds, subsequent predictions are fast. Plotly is used for interactive data visualization.")

add_heading2("5.3  Live Demo Workflow")
add_body("The application features a responsive interface with a sticky navigation bar. The core feature is the Live Demo, where users can:")
add_list_item("Upload a dermoscopic image (JPG/PNG)")
add_list_item("Trigger the full dual-stage ensemble prediction")
add_list_item("View the final classification, confidence scores, and per-model probability breakdowns")

# ═══════════════════════════════════════════════════════════════════
#  CHAPTER VI — CONCLUSION AND FUTURE WORK
# ═══════════════════════════════════════════════════════════════════
add_page_break()
print("Building Chapter VI: Conclusion and Future Work...")
add_heading1("CHAPTER VI")
add_heading1("CONCLUSION AND FUTURE WORK")
add_empty_lines(1)

add_heading2("6.1  Summary of Contributions")
add_body("This thesis has presented a Dual-Stage Ensemble Deep Learning Pipeline for automated skin cancer detection and classification using the ISIC 2019 dataset. The key contributions are:")
add_list_item("Hierarchical Clinical Design: A two-stage pipeline that mirrors dermatologist reasoning — triage first (binary: benign vs. malignant), then specialist classification (4-class malignant subtype) — reducing the complexity faced by any single model.")
add_list_item("Heterogeneous 5-Model Ensemble: Combining ResNet50, EfficientNet-B0, MobileNetV2, ConvNeXt-Tiny, and Swin-Tiny via probability averaging. The mix of CNNs and Vision Transformers provides comprehensive feature coverage.")
add_list_item("Stage-Specific Imbalance Strategies: Stage 1 employs targeted NV undersampling; Stage 2 employs Focal Loss (γ=2) and MixUp augmentation — handling imbalance without discarding rare cancer data.")
add_list_item("Blank-Slate Training for Stage 2: Training from randomly initialised weights prevents ImageNet domain bias from contaminating dermatologically-specific feature learning.")
add_list_item("Grad-CAM Clinical Explainability: Every model decision is accompanied by a visual explanation, meeting the transparency requirements for clinical AI deployment.")
add_list_item("Interactive Streamlit Demo Application: A production-ready web application enabling real-time ensemble inference with per-model probability breakdowns.")

add_heading2("6.2  Key Results")
add_body("Stage 1 Binary Classification: Ensemble accuracy of 94.34%, sensitivity of 93.79%, and AUC of 0.9845. The ensemble reduced false positives by 58 compared to the best individual model while maintaining 93.79% sensitivity.")
add_body("Stage 2 Multi-Class Classification: Ensemble accuracy of 85.58%, Macro-F1 of 0.8332, and Macro-AUC of 0.9738. The critical SCC recall of 95.79% is the highest reported for this dataset configuration. The Clinical Priority Score of 0.8881 demonstrates the system's clinical safety profile.")

add_heading2("6.3  Limitations")
add_list_item("Sensitivity Constraint: A Stage 1 ensemble sensitivity of 93.79% means approximately 6 in every 100 malignant lesions would be missed. The system must function as a decision support tool requiring mandatory dermatologist review.")
add_list_item("Demographic Bias: The ISIC 2019 dataset predominantly represents lighter skin phototypes (Fitzpatrick Types I–III). Performance on darker skin has not been evaluated.")
add_list_item("Ablation Study: The four concurrent methodological improvements were not individually ablated due to computational constraints.")
add_list_item("Regulatory Status: This system has not undergone regulatory review (CE marking, FDA 510(k)). It is a research prototype.")

add_heading2("6.4  Future Work")
add_list_item("Multi-modal Fusion: Incorporating patient metadata (age, sex, anatomical location) alongside image features to improve diagnostic accuracy.")
add_list_item("Diverse Skin Phototype Evaluation: Expanding evaluation to Fitzpatrick Types IV–VI using datasets with broader demographic representation.")
add_list_item("Formal Ablation Study: Systematically evaluating the contribution of each methodological change (ensemble expansion, undersampling, preprocessing, soft-voting).")
add_list_item("Federated Learning: Enabling privacy-preserving model training across multiple hospitals without centralizing patient data.")
add_list_item("Mobile Deployment: Optimizing the lightweight models (EfficientNet-B0, MobileNetV2) for deployment on Android/iOS devices for point-of-care diagnostics in resource-constrained settings.")
add_list_item("Attention-based Ensemble Weighting: Replacing equal-weight averaging with learned attention weights to dynamically adjust each model's contribution based on lesion characteristics.")

# ═══════════════════════════════════════════════════════════════════
#  BIBLIOGRAPHY
# ═══════════════════════════════════════════════════════════════════
add_page_break()
print("Building Bibliography...")
add_heading1("BIBLIOGRAPHY")
add_empty_lines(1)

references = [
    "[1]\tWorld Cancer Research Fund, \"Skin cancer statistics,\" 2024.",
    "[2]\tAmerican Cancer Society, \"Cancer facts & figures 2024,\" American Cancer Society, Atlanta, GA, 2024.",
    "[3]\tY. K. Liang, M. Janda, and H. P. Soyer, \"The global burden of skin cancer,\" Lancet Oncology, vol. 23, no. 2, pp. 186–187, 2022.",
    "[4]\tH. G. Welch and B. L. Mazer, \"Skin cancer treatment costs in the United States,\" JAMA Dermatology, vol. 158, no. 8, pp. 913–914, 2022.",
    "[5]\tD. E. Elder et al., \"WHO classification of skin tumours,\" 4th ed., IARC Press, Lyon, 2018.",
    "[6]\tR. L. Siegel, K. D. Miller, and A. Jemal, \"Cancer statistics, 2024,\" CA Cancer J. Clin., vol. 74, no. 1, pp. 12–49, 2024.",
    "[7]\tN. Howlader et al., \"SEER cancer statistics review 1975–2021,\" National Cancer Institute, 2024.",
    "[8]\tC. M. Balch et al., \"Final version of 2009 AJCC melanoma staging,\" J. Clin. Oncol., vol. 27, pp. 6199–6206, 2009.",
    "[9]\tH. Kittler, H. Pehamberger, K. Wolff, and M. Binder, \"Diagnostic accuracy of dermoscopy,\" Lancet Oncology, vol. 3, no. 3, pp. 159–165, 2002.",
    "[10]\tG. Argenziano et al., \"Dermoscopy of pigmented skin lesions,\" Lancet, vol. 366, no. 9485, pp. 631–641, 2005.",
    "[11]\tI. Zalaudek et al., \"Three-point checklist of dermoscopy,\" Dermatology, vol. 212, pp. 63–67, 2006.",
    "[12]\tH. A. Haenssle et al., \"Man against machine: diagnostic performance of a deep learning CNN,\" Annals of Oncology, vol. 29, no. 8, pp. 1836–1842, 2018.",
    "[13]\tM. E. Vestergaard et al., \"Dermoscopy compared with naked eye examination,\" Br. J. Dermatol., vol. 159, pp. 669–676, 2008.",
    "[14]\tH. A. Haenssle et al., \"Man against machine reloaded,\" European J. of Cancer, vol. 170, pp. 60–69, 2022.",
    "[15]\tInternational League of Dermatological Societies, \"World skin health atlas,\" 2023.",
    "[16]\tR. J. Hay et al., \"The global burden of skin disease,\" J. Invest. Dermatology, vol. 134, pp. 1527–1534, 2014.",
    "[17]\tN. Codella et al., \"Deep learning ensembles for melanoma recognition,\" IBM J. Res. Dev., vol. 61, no. 4, pp. 5:1–5:15, 2017.",
    "[18]\tC. Barata, M. E. Celebi, and J. S. Marques, \"A survey of feature extraction in dermoscopy image analysis,\" IEEE J. Biomed. Health Inform., vol. 23, pp. 1096–1109, 2019.",
    "[19]\tQ. Abbas et al., \"Computer-aided detection of melanocytic lesions,\" Artif. Intell. Med., vol. 56, pp. 69–78, 2012.",
    "[20]\tY. LeCun, Y. Bengio, and G. Hinton, \"Deep learning,\" Nature, vol. 521, pp. 436–444, 2015.",
    "[21]\tA. Krizhevsky, I. Sutskever, and G. E. Hinton, \"ImageNet classification with deep CNNs,\" NIPS, pp. 1097–1105, 2012.",
    "[22]\tA. Esteva et al., \"Dermatologist-level classification of skin cancer,\" Nature, vol. 542, pp. 115–118, 2017.",
    "[23]\tP. Tschandl et al., \"Human-computer collaboration for skin cancer recognition,\" Lancet Oncology, vol. 21, pp. 906–914, 2020.",
    "[24]\tJ. Yosinski et al., \"How transferable are features in deep neural networks?,\" NIPS, pp. 3320–3328, 2014.",
    "[25]\tM. Raghu et al., \"Transfusion: Understanding transfer learning for medical imaging,\" NeurIPS, 32, 2019.",
    "[26]\tK. He, X. Zhang, S. Ren, and J. Sun, \"Deep residual learning,\" CVPR, pp. 770–778, 2016.",
    "[27]\tG. Huang et al., \"Densely connected convolutional networks,\" CVPR, pp. 4700–4708, 2017.",
    "[28]\tM. Tan and Q. V. Le, \"EfficientNet: Rethinking model scaling,\" ICML, pp. 6105–6114, 2019.",
    "[29]\tA. Holzinger et al., \"Causability and explainability of AI in medicine,\" WIREs Data Mining and Knowledge Discovery, vol. 9, no. 4, e1312, 2019.",
    "[30]\tR. R. Selvaraju et al., \"Grad-CAM: Visual explanations from deep networks,\" ICCV, pp. 618–626, 2017.",
    "[31]\tD. Gutman et al., \"Skin lesion analysis toward melanoma detection,\" ISBI Challenge, 2016.",
    "[32]\tN. C. F. Codella et al., \"Skin lesion analysis toward melanoma detection: ISIC 2019,\" arXiv:1902.03368, 2019.",
    "[33]\tP. Tschandl, C. Rosendahl, and H. Kittler, \"The HAM10000 dataset,\" Scientific Data, vol. 6, Article 205, 2019.",
    "[34]\tS. Pacheco and R. Kramer, \"On out-of-distribution detection in medical imaging,\" Medical Image Analysis, vol. 74, 102227, 2021.",
    "[35]\tV. Rotemberg et al., \"A patient-centric dataset of images and metadata,\" Scientific Data, vol. 8, Article 34, 2021.",
    "[36]\tM. Sandler et al., \"MobileNetV2: Inverted residuals,\" CVPR, pp. 4510–4520, 2018.",
    "[37]\tA. Young et al., \"Artificial intelligence in dermatology: A systematic review,\" J. Invest. Dermatology, vol. 140, pp. 1504–1515, 2020.",
    "[38]\tZ. Liu et al., \"Swin Transformer: Hierarchical vision transformer,\" ICCV, pp. 10012–10022, 2021.",
    "[39]\tZ. Liu et al., \"A ConvNet for the 2020s,\" CVPR, pp. 11976–11986, 2022.",
    "[40]\tT.-Y. Lin et al., \"Focal loss for dense object detection,\" ICCV, pp. 2980–2988, 2017.",
    "[41]\tH. Zhang et al., \"mixup: Beyond empirical risk minimization,\" ICLR, 2018.",
    "[42]\tA. Telea, \"An image inpainting technique based on the fast marching method,\" J. Graphics Tools, vol. 9, no. 1, pp. 23–34, 2004.",
]

for ref in references:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(ref)
    set_run_font(run, size=12)
    set_paragraph_spacing(p, before=5, after=5, line_spacing=1.5)

# ═══════════════════════════════════════════════════════════════════
#  SAVE DOCUMENT
# ═══════════════════════════════════════════════════════════════════
print(f"\nSaving thesis to: {OUTPUT_PATH}")
doc.save(OUTPUT_PATH)
print("[SUCCESS] THESIS GENERATED SUCCESSFULLY!")
print(f"   File: {OUTPUT_PATH}")

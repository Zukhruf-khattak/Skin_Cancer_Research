import docx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def extract_and_apply_page_numbers(filename):
    doc = docx.Document(filename)
    
    print("=== Extracting Page Numbers ===")
    
    # 1. First, let's map out exactly which page every paragraph is on.
    # In Word OOXML, a page break can be explicit (w:br type="page") 
    # or rendered (w:lastRenderedPageBreak).
    # Also, a section break (w:sectPr) often starts a new page.
    
    # Let's track the page number. Front matter uses Roman, but the user said 
    # "page number from chapter 1 introduction". 
    # Usually Chapter 1 is page 1. Let's find Chapter 1 and reset the counter there!
    
    current_physical_page = 1
    page_mapping = {} # paragraph index -> page string
    
    # Find chapter 1 index to reset page counter
    ch1_idx = -1
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().upper() == 'CHAPTER I':
            ch1_idx = i
            break
            
    is_front_matter = True
    display_page = 1
    roman_numerals = ["i", "ii", "iii", "iv", "v", "vi", "vii", "viii", "ix", "x", "xi", "xii", "xiii"]
    
    for i, p in enumerate(doc.paragraphs):
        if i == ch1_idx:
            is_front_matter = False
            display_page = 1 # Reset to 1 at Chapter 1
            
        # Count breaks IN this paragraph
        breaks = p._p.findall('.//' + qn('w:br'))
        rendered = p._p.findall('.//' + qn('w:lastRenderedPageBreak'))
        
        # Some paragraphs have multiple rendered breaks if they span multiple pages.
        # But a heading is usually just on one page.
        
        # If there are breaks BEFORE the text in this paragraph, the text is on the NEW page.
        # It's tricky to parse exact position, but generally if a rendered break exists, 
        # it means a page boundary was crossed in this paragraph.
        
        if is_front_matter:
            page_str = roman_numerals[display_page - 1] if display_page <= len(roman_numerals) else str(display_page)
        else:
            page_str = str(display_page)
            
        page_mapping[i] = page_str
        
        # Now increment for the NEXT paragraph if we hit a break
        num_breaks = sum(1 for br in breaks if br.get(qn('w:type')) == 'page')
        num_breaks += len(rendered)
        
        # Check section break
        pPr = p._p.find(qn('w:pPr'))
        if pPr is not None and pPr.find(qn('w:sectPr')) is not None:
            num_breaks += 1
            
        display_page += num_breaks

    # 2. Map headings to page numbers
    toc_data = []
    toc_indices = []
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        style = para.style.name
        
        if not text:
            continue
            
        level = None
        if re.match(r'^CHAPTER\s+(I|II|III|IV|V|VI|VII|VIII|IX|X)\b', text, re.IGNORECASE):
            level = 1
        elif text.upper() == 'REFERENCES':
            level = 1
        elif 'Heading 2' in style:
            level = 2
        elif 'Heading 3' in style:
            level = 3
            
        if level is not None:
            # We must NOT include the TOC entries themselves!
            # The TOC entries are in the front matter, they will have level=1, 2, 3 etc. 
            # Wait, the TOC entries I generated don't have Heading styles, they have Normal.
            # Except CHAPTER I, but I didn't style the TOC entries as Heading 1!
            # Wait, in the manual TOC, what style did they get? Normal!
            # But the actual headings have 'Heading 1', 'Heading 2'.
            # Wait, my TOC entries are just text. 
            # But wait, what about the TOC entries themselves? 
            # I can skip paras before ch1_idx. But front-matter headings like ABSTRACT? 
            # They don't have Heading styles usually, except ABSTRACT might be Heading 1.
            # Let's just track the bookmark anchors if they exist!
            
            # Find bookmarkStart with name starting with TOC_BM_
            bks = para._p.findall('.//' + qn('w:bookmarkStart'))
            for bk in bks:
                if bk.get(qn('w:name'), '').startswith('TOC_BM_'):
                    toc_data.append((text, level, page_mapping[i]))
                    toc_indices.append(i)
                    print(f"Matched {text[:30]} -> Page {page_mapping[i]}")
                    break

    # Now we have toc_data: list of (text, level, page_str)
    # We will modify the TOC entries IN PLACE in the document!
    
    print("\n=== Updating TOC Entries In-Place ===")
    updated_count = 0
    for p in doc.paragraphs:
        hls = p._p.findall('.//' + qn('w:hyperlink'))
        for hl in hls:
            anchor = hl.get(qn('w:anchor'), '')
            if anchor.startswith('TOC_BM_'):
                # Extract the index from the anchor (e.g., TOC_BM_0)
                try:
                    idx = int(anchor.split('_')[-1])
                    if idx < len(toc_data):
                        heading_text, level, page_str = toc_data[idx]
                        
                        # We need to add the dot leader and page number to this paragraph.
                        # Since we want it clickable, we can put it inside the hyperlink, 
                        # or outside. The user wants hardcoded numbers so we don't need fields!
                        # We'll just put it OUTSIDE the hyperlink so it doesn't get underlined in blue.
                        
                        # First, make sure we haven't already added it (avoid duplicates)
                        existing_tabs = p._p.findall('.//' + qn('w:tab'))
                        if len(existing_tabs) == 0:
                            # 1. Ensure tab stop is in pPr
                            pPr = p._p.find(qn('w:pPr'))
                            if pPr is None:
                                pPr = OxmlElement('w:pPr')
                                p._p.insert(0, pPr)
                            tabs = pPr.find(qn('w:tabs'))
                            if tabs is None:
                                tabs = OxmlElement('w:tabs')
                                pPr.append(tabs)
                            tab_stop = OxmlElement('w:tab')
                            tab_stop.set(qn('w:val'), 'right')
                            tab_stop.set(qn('w:leader'), 'dot')
                            tab_stop.set(qn('w:pos'), '9000') # ~6.25 inches
                            tabs.append(tab_stop)
                            
                            # 2. Add the literal text run for page number
                            r = OxmlElement('w:r')
                            rPr_r = OxmlElement('w:rPr')
                            rFonts = OxmlElement('w:rFonts')
                            rFonts.set(qn('w:ascii'), 'Times New Roman')
                            rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                            rPr_r.append(rFonts)
                            sz = OxmlElement('w:sz')
                            font_size_hpc = {1: 24, 2: 24, 3: 24}.get(level, 24)
                            sz.set(qn('w:val'), str(font_size_hpc))
                            rPr_r.append(sz)
                            # Make page number black
                            col = OxmlElement('w:color')
                            col.set(qn('w:val'), '000000')
                            rPr_r.append(col)
                            
                            if level == 1:
                                rPr_r.append(OxmlElement('w:b'))
                            
                            r.append(rPr_r)
                            r.append(OxmlElement('w:tab'))
                            
                            t = OxmlElement('w:t')
                            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                            t.text = page_str
                            r.append(t)
                            
                            p._p.append(r)
                            updated_count += 1
                except Exception as e:
                    print(f"Error processing anchor {anchor}: {e}")
                    
    print(f"Updated {updated_count} TOC entries with hardcoded page numbers!")
    
    out_file = r'd:\RESEARCH PNGS\zukhruftheisis_FINAL_TOC_Numbered.docx'
    doc.save(out_file)
    print(f"Saved to {out_file}")

extract_and_apply_page_numbers(r'd:\RESEARCH PNGS\zukhruftheisis_FINAL_TOC.docx')

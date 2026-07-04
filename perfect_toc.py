import docx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_perfect_toc(in_file, out_file):
    doc = docx.Document(in_file)
    
    # 1. Map paragraphs to bookmarks
    bkmk_to_para = {}
    for i, p in enumerate(doc.paragraphs):
        for bk in p._p.findall('.//' + qn('w:bookmarkStart')):
            name = bk.get(qn('w:name'), '')
            if name.startswith('TOC_BM_'):
                bkmk_to_para[name] = i
                
    start_idx = bkmk_to_para['TOC_BM_0']
    
    # 2. Count pages starting from start_idx
    current_page = 1
    mapping = {}
    
    # The first rendered break at or immediately after start_idx is the boundary 
    # between the TOC and Chapter 1. We must ignore it.
    ignored_first_break = False
    
    for i in range(start_idx, len(doc.paragraphs)):
        p = doc.paragraphs[i]
        
        # Check if this paragraph has a heading bookmark
        for bk in p._p.findall('.//' + qn('w:bookmarkStart')):
            name = bk.get(qn('w:name'), '')
            if name.startswith('TOC_BM_'):
                mapping[name] = str(current_page)
                
        # Count rendered breaks
        rendered = p._p.findall('.//' + qn('w:lastRenderedPageBreak'))
        num_breaks = len(rendered)
        
        for _ in range(num_breaks):
            if not ignored_first_break:
                ignored_first_break = True # ignore the boundary before Chapter 1
            else:
                current_page += 1

    # 3. Update the TOC IN PLACE
    updated = 0
    for i, p in enumerate(doc.paragraphs):
        hls = p._p.findall('.//' + qn('w:hyperlink'))
        for hl in hls:
            anchor = hl.get(qn('w:anchor'), '')
            if anchor.startswith('TOC_BM_'):
                page_str = mapping.get(anchor, "1")
                
                # Check if it already has tabs (meaning we injected it already)
                existing_tabs = p._p.findall('.//' + qn('w:tab'))
                if len(existing_tabs) == 0:
                    # Inject from scratch
                    pPr = p._p.find(qn('w:pPr'))
                    if pPr is None:
                        pPr = OxmlElement('w:pPr')
                        p._p.insert(0, pPr)
                    tabs = pPr.find(qn('w:tabs'))
                    if tabs is None:
                        tabs = OxmlElement('w:tabs')
                        pPr.append(tabs)
                    tab_stop = OxmlElement('w:tab')
                    tab_stop.set(qn('w:val'), 'right')
                    tab_stop.set(qn('w:leader'), 'dot')
                    tab_stop.set(qn('w:pos'), '9000') 
                    tabs.append(tab_stop)
                    
                    r = OxmlElement('w:r')
                    rPr_r = OxmlElement('w:rPr')
                    rFonts = OxmlElement('w:rFonts')
                    rFonts.set(qn('w:ascii'), 'Times New Roman')
                    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                    rPr_r.append(rFonts)
                    sz = OxmlElement('w:sz')
                    
                    level = 2
                    text_run = hl.find('.//' + qn('w:t'))
                    if text_run is not None and text_run.text and text_run.text.strip().upper().startswith("CHAPTER"):
                        level = 1
                    elif text_run is not None and text_run.text and text_run.text.strip().upper() == "REFERENCES":
                        level = 1
                        
                    font_size = {1: 24, 2: 24, 3: 24}.get(level, 24)
                    sz.set(qn('w:val'), str(font_size))
                    rPr_r.append(sz)
                    
                    col = OxmlElement('w:color')
                    col.set(qn('w:val'), '000000')
                    rPr_r.append(col)
                    
                    if level == 1:
                        rPr_r.append(OxmlElement('w:b'))
                    
                    r.append(rPr_r)
                    r.append(OxmlElement('w:tab'))
                    
                    t = OxmlElement('w:t')
                    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                    t.text = page_str
                    r.append(t)
                    
                    p._p.append(r)
                    updated += 1
                else:
                    # Update existing page number
                    runs = p._p.findall('.//' + qn('w:r'))
                    if runs:
                        last_r = runs[-1]
                        t_nodes = last_r.findall('.//' + qn('w:t'))
                        if t_nodes:
                            t_nodes[-1].text = page_str
                            updated += 1

    print(f"Updated {updated} TOC entries with PERFECT page numbers!")
    doc.save(out_file)
    print(f"Saved to {out_file}")

create_perfect_toc(r'd:\RESEARCH PNGS\zukhruftheisis_FINAL_TOC.docx', r'd:\RESEARCH PNGS\zukhruftheisis_FINAL_TOC_Numbered.docx')

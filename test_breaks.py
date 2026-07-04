import docx
from docx.oxml.ns import qn

def test_rendered_breaks(filename):
    doc = docx.Document(filename)
    
    start_idx = 149
    current_page = 1
    
    for i in range(start_idx, len(doc.paragraphs)):
        p = doc.paragraphs[i]
        
        rendered = p._p.findall('.//' + qn('w:lastRenderedPageBreak'))
        if rendered:
            current_page += len(rendered)
            
        for bk in p._p.findall('.//' + qn('w:bookmarkStart')):
            name = bk.get(qn('w:name'), '')
            if name.startswith('TOC_BM_'):
                if '0' in name or '10' in name or '101' in name:
                    print(f"Heading {name} ('{p.text.strip()[:30]}') -> Page {current_page}")

test_rendered_breaks(r'd:\RESEARCH PNGS\zukhruftheisis_FINAL_TOC.docx')

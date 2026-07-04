"""
Comprehensive thesis modifier for zukhruftheisis.docx:
1. Shorten Introduction - trim verbose background sections 1.1-1.7
2. Shorten Literature Review - trim 2.11 Recent Studies (reduce app count)
3. Fix ALL fonts: Times New Roman throughout
   - Body text: 12pt
   - Subheadings (Heading 3): 12pt bold black
   - Main headings (Heading 2): 14pt bold black  
   - Chapter title / Title: 16pt bold black
"""

import docx
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy
import re

# =====================================================================
# LOAD
# =====================================================================
doc = Document(r'd:\RESEARCH PNGS\zukhruftheisis.docx')
paras = doc.paragraphs

# =====================================================================
# HELPER: Apply font to a single paragraph
# =====================================================================
def apply_font(para, name='Times New Roman', size_pt=12, bold=None, color=(0,0,0)):
    """Apply font to all runs in a paragraph."""
    for run in para.runs:
        run.font.name = name
        run.font.size = Pt(size_pt)
        if bold is not None:
            run.font.bold = bold
        if color is not None:
            run.font.color.rgb = RGBColor(*color)
        # Ensure theme fonts also updated in XML
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:ascii'), name)
        rFonts.set(qn('w:hAnsi'), name)
        rFonts.set(qn('w:cs'), name)
    
    # Also fix paragraph-level font in pPr > rPr
    pPr = para._p.find(qn('w:pPr'))
    if pPr is not None:
        pRPr = pPr.find(qn('w:rPr'))
        if pRPr is not None:
            rFonts = pRPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                pRPr.insert(0, rFonts)
            rFonts.set(qn('w:ascii'), name)
            rFonts.set(qn('w:hAnsi'), name)
            rFonts.set(qn('w:cs'), name)


# =====================================================================
# HELPER: Remove a paragraph from document
# =====================================================================
def remove_paragraph(para):
    """Remove a paragraph element from the document."""
    p = para._element
    p.getparent().remove(p)


# =====================================================================
# HELPER: Replace a paragraph's text (preserving runs)
# =====================================================================
def set_para_text(para, new_text):
    """Clear all runs and set new text in first run."""
    for run in para.runs:
        run.text = ''
    if para.runs:
        para.runs[0].text = new_text
    else:
        para.add_run(new_text)


# =====================================================================
# STEP 1: FIX FONT FORMATTING
# =====================================================================
print("Step 1: Fixing font formatting...")

for para in doc.paragraphs:
    text = para.text.strip()
    style = para.style.name

    # Determine category
    is_chapter = (
        re.match(r'^CHAPTER\s+(I|II|III|IV|V|VI|VII|VIII|IX|X)', text, re.IGNORECASE) or
        (text.isupper() and len(text.split()) <= 5 and len(text) > 2 and style in ['Normal', 'No Spacing', 'Title']) or
        style in ['Title', 'Subtitle'] or
        'Heading 1' in style
    )
    is_heading2 = 'Heading 2' in style and not is_chapter
    is_heading3 = 'Heading 3' in style

    if is_chapter:
        apply_font(para, size_pt=16, bold=True)
    elif is_heading2:
        apply_font(para, size_pt=14, bold=True)
    elif is_heading3:
        apply_font(para, size_pt=12, bold=True)
    else:
        apply_font(para, size_pt=12, bold=None)

print("  Font formatting complete.")

# =====================================================================
# STEP 2: SHORTEN INTRODUCTION
# =====================================================================
# Reload paragraph references (doc.paragraphs is a live list but we 
# need to collect indices before removing)
print("\nStep 2: Shortening Introduction...")

# We will REPLACE verbose body paragraphs with condensed single paragraphs.
# Strategy: for each background section (1.1-1.7), keep only the heading
# and replace multiple body paras with 1 concise paragraph.

# New condensed text for each section body
condensed_intro = {
    '1.1 Background': (
        "Skin cancer is one of the most prevalent malignancies worldwide. Early and accurate detection is critical "
        "because survival rates drop significantly with delayed diagnosis. Dermoscopy provides magnified skin surface "
        "visualization, but manual analysis is time-consuming and subject to inter-observer variability. Computer-aided "
        "diagnosis (CAD) systems powered by deep learning offer a scalable approach to assist clinicians in distinguishing "
        "benign from malignant lesions with high accuracy."
    ),
    '1.2 Classification': (
        "Skin lesions are broadly classified as benign or malignant. This study focuses on the diagnostic classes in the "
        "ISIC 2019 dataset. Malignant and pre-malignant categories include Melanoma (MEL), Basal Cell Carcinoma (BCC), "
        "Squamous Cell Carcinoma (SCC), and Actinic Keratosis (AK). Benign categories include Melanocytic Nevus (NV), "
        "Benign Keratosis-like Lesions (BKL), Dermatofibroma (DF), and Vascular Lesions (VASC). The visual similarity "
        "between benign and malignant lesions makes automated classification particularly challenging and clinically important."
    ),
    '1.3 Clinical Diagnosis': (
        "Clinical diagnosis of skin cancer relies on visual inspection, dermoscopy, and biopsy confirmation. Dermoscopic "
        "analysis requires specialized training and is prone to inter-observer variability. Access to experienced "
        "dermatologists is limited in many regions, motivating the development of automated diagnostic support systems "
        "that can provide consistent and accessible pre-screening."
    ),
    '1.4 CAD Systems': (
        "Computer-aided diagnosis (CAD) systems have evolved from handcrafted-feature pipelines to end-to-end deep "
        "learning frameworks. Early CAD systems relied on manual feature engineering, including color histograms, "
        "texture descriptors, and shape features. Deep learning overcame these limitations by learning discriminative "
        "features directly from images, achieving dermatologist-level performance on benchmark datasets."
    ),
    '1.5 Deep Learning': (
        "Deep learning, particularly convolutional neural networks (CNNs), has transformed medical image analysis. "
        "Landmark studies demonstrated that CNN-based models can match or exceed dermatologist accuracy on dermoscopic "
        "classification tasks. More recently, Vision Transformers (ViT) and hybrid architectures have extended these "
        "capabilities by capturing long-range spatial dependencies alongside local texture features."
    ),
    '1.6 Transfer Learning': (
        "Transfer learning enables deep networks to leverage features learned from large-scale datasets (e.g., ImageNet) "
        "for tasks with limited labeled data. In skin lesion classification, fine-tuning pre-trained models consistently "
        "outperforms training from scratch, particularly when labeled clinical data is scarce or imbalanced. "
        "This study applies transfer learning as the primary training strategy for both stages of the proposed pipeline."
    ),
    '1.7 XAI': (
        "Explainable AI (XAI) is essential in clinical decision support, where black-box predictions are insufficient "
        "for clinician trust. Gradient-weighted Class Activation Mapping (Grad-CAM) generates saliency maps that "
        "highlight which image regions most influenced a model's decision. This study integrates Grad-CAM to verify "
        "that the proposed models attend to dermoscopically meaningful lesion areas."
    ),
}

# Map heading text prefixes to condensed bodies
section_map = {
    '1.1': (51, 57),
    '1.2': (57, 71),
    '1.3': (71, 76),
    '1.4': (76, 81),
    '1.5': (81, 86),
    '1.6': (86, 91),
    '1.7': (91, 96),
}

# We need to remove body paragraphs and replace with single condensed one.
# We do this by:
# 1. Mark all body paras for deletion
# 2. Set text of first body para to condensed text
# 3. Delete the rest

paras_to_delete = []
paras_to_replace = {}  # para_index: new_text

# Map section number prefix to condensed text
condensed_by_number = {
    '1.1': condensed_intro['1.1 Background'],
    '1.2': condensed_intro['1.2 Classification'],
    '1.3': condensed_intro['1.3 Clinical Diagnosis'],
    '1.4': condensed_intro['1.4 CAD Systems'],
    '1.5': condensed_intro['1.5 Deep Learning'],
    '1.6': condensed_intro['1.6 Transfer Learning'],
    '1.7': condensed_intro['1.7 XAI'],
}

for key, (h_idx, next_h_idx) in section_map.items():
    body_start = h_idx + 1
    body_end = next_h_idx
    condensed_key = key
    
    # Keep first body para, replace its text; mark rest for deletion
    first_body = body_start
    if first_body < body_end:
        paras_to_replace[first_body] = condensed_by_number[condensed_key]
        for del_idx in range(first_body + 1, body_end):
            paras_to_delete.append(del_idx)
        print(f"  Condensing: '{paras[h_idx].text}' => keeping para {first_body}, deleting {first_body+1}:{body_end}")
    else:
        print(f"  Skipping (no body paras): '{paras[h_idx].text}'")

# Apply replacements BEFORE deletions (indices will shift during deletion)
for idx, text in paras_to_replace.items():
    para = doc.paragraphs[idx]
    # Clear all runs
    for run in para.runs:
        run.text = ''
    if para.runs:
        para.runs[0].text = text
    else:
        run = para.add_run(text)
    # Fix font on new content
    apply_font(para, size_pt=12, bold=False)

# Sort deletion list descending so indices remain valid
paras_to_delete_sorted = sorted(set(paras_to_delete), reverse=True)
print(f"\n  Deleting {len(paras_to_delete_sorted)} paragraphs from introduction...")

for idx in paras_to_delete_sorted:
    para = doc.paragraphs[idx]
    # Only delete if it's a normal body paragraph (not a heading)
    if 'Heading' not in para.style.name:
        p = para._element
        p.getparent().remove(p)

print("  Introduction shortened.")

# =====================================================================
# STEP 3: SHORTEN LITERATURE REVIEW - Section 2.11 (Recent Studies)
# =====================================================================
print("\nStep 3: Shortening Literature Review section 2.11...")

# After deleting intro paras, indices have shifted. Re-find section 2.11.
def find_para_by_text(search, start=0):
    for i in range(start, len(doc.paragraphs)):
        if search.lower() in doc.paragraphs[i].text.lower():
            return i
    return -1

sec211_idx = find_para_by_text('2.11 Review of Recent Studies')
sec212_idx = find_para_by_text('2.12 Research Gaps')

print(f"  Section 2.11 found at paragraph index: {sec211_idx}")
print(f"  Section 2.12 found at paragraph index: {sec212_idx}")

if sec211_idx != -1 and sec212_idx != -1:
    # New condensed version of 2.11 (1-2 paragraphs, 3 apps instead of many)
    condensed_211_text = (
        "Recent studies confirm that skin lesion classification research has converged toward transfer learning, "
        "ensemble approaches, and explainable AI. Aljohani and Turki demonstrated CNN-based melanoma classification "
        "achieving strong sensitivity on dermoscopic benchmarks [32]. Tschandl et al. showed that deep learning "
        "ensembles outperform individual dermatologists when evaluated on the HAM10000 dataset [33]. Esteva et al. "
        "established that CNNs trained on diverse dermoscopic images can reach specialist-level accuracy for melanoma "
        "versus benign lesion discrimination [34]. These findings collectively support the use of ensemble deep "
        "learning with transfer learning as the foundation for the proposed dual-stage pipeline. A key gap across "
        "these studies is the lack of integrated multi-stage classification and clinical explainability, which this "
        "research directly addresses."
    )

    # Replace the body of section 2.11
    body_start = sec211_idx + 1
    body_end = sec212_idx

    # Replace first body para with condensed text
    if body_start < body_end:
        first_para = doc.paragraphs[body_start]
        for run in first_para.runs:
            run.text = ''
        if first_para.runs:
            first_para.runs[0].text = condensed_211_text
        else:
            first_para.add_run(condensed_211_text)
        apply_font(first_para, size_pt=12, bold=False)

        # Delete the rest (descending order)
        to_delete = list(range(body_start + 1, body_end))
        print(f"  Deleting {len(to_delete)} paragraphs from section 2.11 body...")
        for idx in sorted(to_delete, reverse=True):
            para = doc.paragraphs[idx]
            if 'Heading' not in para.style.name:
                p = para._element
                p.getparent().remove(p)

    print("  Literature Review section 2.11 shortened.")

# =====================================================================
# STEP 4: ALSO TRIM SECTION 2.10 (Public Datasets) - remove empty paras
# =====================================================================
print("\nStep 4: Cleaning empty paragraphs in Literature Review...")
# Remove blank/empty body paragraphs (from spacing issues)
empty_paras = []
for i, para in enumerate(doc.paragraphs):
    if not para.text.strip() and para.style.name in ['Normal', 'No Spacing']:
        empty_paras.append(i)

# Only remove in groups (to not strip intentional single line breaks)
# Remove stretches of 2+ consecutive empty paragraphs
to_remove = []
for i in range(len(empty_paras) - 1):
    if empty_paras[i+1] - empty_paras[i] == 1:
        to_remove.append(empty_paras[i])

print(f"  Removing {len(to_remove)} excess blank paragraphs...")
for idx in sorted(set(to_remove), reverse=True):
    try:
        para = doc.paragraphs[idx]
        p = para._element
        p.getparent().remove(p)
    except:
        pass

# =====================================================================
# SAVE
# =====================================================================
output_path = r'd:\RESEARCH PNGS\zukhruftheisis_revised.docx'
doc.save(output_path)
print("\nDONE! Saved to: " + output_path)
print("Changes made:")
print("  1. Font: Times New Roman throughout")
print("  2. Body text: 12pt")
print("  3. Subheadings (Heading 3): 12pt bold black")
print("  4. Main Headings (Heading 2): 14pt bold black")
print("  5. Chapter / Title text: 16pt bold black")
print("  6. Introduction sections 1.1-1.7 condensed to 1 paragraph each")
print("  7. Literature Review section 2.11 condensed (3 key studies retained)")
print("  8. Excess blank paragraphs removed")

import docx
from docx.shared import Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE

def update_thesis_font(filepath, output_filepath):
    doc = docx.Document(filepath)
    font_name = 'Times New Roman'
    
    # Update Styles (Normal, Headings, etc.)
    for style in doc.styles:
        if style.type == WD_STYLE_TYPE.PARAGRAPH:
            if style.font:
                style.font.name = font_name
                    
    # Update direct formatting in all paragraphs and runs
    for para in doc.paragraphs:
        for run in para.runs:
            run.font.name = font_name
            
    # Also update tables if any
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.name = font_name

    doc.save(output_filepath)
    print(f"Successfully saved updated document to {output_filepath}")

if __name__ == '__main__':
    input_file = 'Thesis_Chapters_1_to_6_Corrected_Final.docx'
    output_file = 'Thesis_Chapters_1_to_6_Corrected_Final.docx'
    update_thesis_font(input_file, output_file)

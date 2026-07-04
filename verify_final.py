from docx import Document
from docx.oxml.ns import qn
import re

doc = Document(r'd:\RESEARCH PNGS\zukhruftheisis_final.docx')

# Check reference count
refs_start = -1
for i, p in enumerate(doc.paragraphs):
    if p.text.strip().upper() == 'REFERENCES':
        refs_start = i
        break

ref_count = 0
for i in range(refs_start+1, len(doc.paragraphs)):
    if doc.paragraphs[i].text.strip():
        ref_count += 1
print('Reference count:', ref_count)

# Check captions font size
cap_check = []
for p in doc.paragraphs:
    t = p.text.strip()
    if re.match(r'^(Table|Figure)\s+[\d\.]+', t, re.IGNORECASE) and p.runs:
        fs = p.runs[0].font.size
        fs_pt = round(fs/12700) if fs else None
        cap_check.append((fs_pt, t[:60]))
print('Sample caption sizes (' + str(len(cap_check)) + ' total):')
for s, t in cap_check[:5]:
    print('  ' + str(s) + 'pt - ' + t)

# Check page numbering sectPr
body = doc.element.body
body_sectPr = body.find(qn('w:sectPr'))
pnt = body_sectPr.find(qn('w:pgNumType'))
if pnt is not None:
    print('\nBody sectPr (Chapter I onwards): fmt=' + str(pnt.get(qn('w:fmt'))) + ', start=' + str(pnt.get(qn('w:start'))))
else:
    print('\nBody sectPr: No pgNumType')

# Find front-matter sectPr
for i, para in enumerate(doc.paragraphs):
    pPr = para._p.find(qn('w:pPr'))
    if pPr is not None:
        s = pPr.find(qn('w:sectPr'))
        if s is not None:
            pnt2 = s.find(qn('w:pgNumType'))
            fmt_val = pnt2.get(qn('w:fmt')) if pnt2 is not None else None
            start_val = pnt2.get(qn('w:start')) if pnt2 is not None else None
            print('Para [' + str(i) + '] sectPr: fmt=' + str(fmt_val) + ', start=' + str(start_val))

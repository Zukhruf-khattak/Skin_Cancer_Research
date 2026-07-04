import docx
from docx.shared import Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_LINE_SPACING
import copy

def update_thesis_formatting(filepath, output_filepath):
    doc = docx.Document(filepath)
    
    black = RGBColor(0, 0, 0)
    font_name = 'Times New Roman'
    
    # -------------------------------------------------------
    # Step 1: Update Paragraph STYLES for headings and normal
    # -------------------------------------------------------
    for style in doc.styles:
        if style.type != WD_STYLE_TYPE.PARAGRAPH:
            continue

        sname = style.name

        # --- Heading 1 = Chapter/Parent Heading: 14pt ---
        if sname == 'Heading 1':
            style.font.name = font_name
            style.font.size = Pt(14)
            style.font.color.rgb = black
            style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        # --- Heading 2/3/4/... = Subheadings: 12pt ---
        elif sname.startswith('Heading '):
            style.font.name = font_name
            style.font.size = Pt(12)
            style.font.color.rgb = black
            style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        # --- Normal body text: 12pt, 1.5 spacing ---
        elif sname in ('Normal', 'Body Text', 'Default Paragraph Font'):
            style.font.name = font_name
            style.font.size = Pt(12)
            style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    # -------------------------------------------------------
    # Step 2: Apply direct formatting to every paragraph
    # -------------------------------------------------------
    for para in doc.paragraphs:
        sname = para.style.name if para.style else ''

        # Determine target size
        if sname == 'Heading 1':
            target_pt = Pt(14)
        elif sname.startswith('Heading '):
            target_pt = Pt(12)
        else:
            target_pt = Pt(12)   # body text

        # Apply 1.5 line spacing at paragraph level
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        # Apply per-run formatting
        for run in para.runs:
            run.font.name = font_name
            run.font.size = target_pt
            if sname.startswith('Heading '):
                run.font.color.rgb = black

    # -------------------------------------------------------
    # Step 3: Also fix text in tables
    # -------------------------------------------------------
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                    for run in para.runs:
                        run.font.name = font_name
                        run.font.size = Pt(12)

    doc.save(output_filepath)
    print("Done! Saved to:", output_filepath)

if __name__ == '__main__':
    f = 'Thesis_Chapters_1_to_6_Corrected_Final.docx'
    update_thesis_formatting(f, f)

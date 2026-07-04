import docx
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import copy

def rebuild_manual_toc(in_file, out_file):
    doc = docx.Document(in_file)
    
    # 1. Remove the existing TOC heading and TOC field code
    paras_to_remove = []
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip().upper()
        if 'TABLE OF CONTENTS' in text:
            paras_to_remove.append(p)
        
        # Check for TOC field code
        has_fld = False
        for instr in p._p.findall('.//' + qn('w:instrText')):
            if instr.text and 'TOC' in instr.text:
                has_fld = True
        if has_fld:
            paras_to_remove.append(p)
            
    for p in paras_to_remove:
        p._element.getparent().remove(p._element)
        
    print(f"Removed {len(paras_to_remove)} old TOC-related paragraphs")

    # 2. Collect all headings and add bookmarks
    toc_entries = []
    bk_counter = 0
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        style = para.style.name
        if not text:
            continue
        
        level = None
        if re.match(r'^CHAPTER\s+(I|II|III|IV|V|VI|VII|VIII|IX|X)\b', text, re.IGNORECASE):
            level = 1
        elif text.upper() == 'REFERENCES':
            level = 1
        elif 'Heading 2' in style:
            level = 2
        elif 'Heading 3' in style:
            level = 3
        
        if level is not None:
            bk_name = f'TOC_BM_{bk_counter}'
            bk_counter += 1
            toc_entries.append((i, text, level, bk_name))

    print(f"Found {len(toc_entries)} headings for TOC")

    # Add bookmarks to heading paragraphs
    for (idx, text, level, bk_name) in toc_entries:
        para = doc.paragraphs[idx]
        # Remove any existing bookmarkStart/End to avoid duplicates
        for bk in para._p.findall('.//' + qn('w:bookmarkStart')):
            if bk.get(qn('w:name'), '').startswith('TOC_BM_'):
                bk.getparent().remove(bk)
        for bk in para._p.findall('.//' + qn('w:bookmarkEnd')):
            bk.getparent().remove(bk)
        
        bk_id = str(9000 + idx)
        bk_start = OxmlElement('w:bookmarkStart')
        bk_start.set(qn('w:id'), bk_id)
        bk_start.set(qn('w:name'), bk_name)
        bk_end = OxmlElement('w:bookmarkEnd')
        bk_end.set(qn('w:id'), bk_id)
        
        para._p.insert(0, bk_start)
        para._p.append(bk_end)

    # 3. Build the TOC elements
    def create_toc_entry_para(text, level, bookmark_name):
        p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        
        indent_twips = {1: 0, 2: 360, 3: 720}.get(level, 0)
        if indent_twips:
            ind = OxmlElement('w:ind')
            ind.set(qn('w:left'), str(indent_twips))
            pPr.append(ind)
        
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:before'), '40')
        spacing.set(qn('w:after'), '40')
        pPr.append(spacing)
        
        p.append(pPr)
        
        font_size_hpc = {1: 24, 2: 22, 3: 20}.get(level, 24)
        is_bold = (level == 1)
        
        def make_rPr():
            rPr = OxmlElement('w:rPr')
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), 'Times New Roman')
            rFonts.set(qn('w:hAnsi'), 'Times New Roman')
            rPr.append(rFonts)
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), str(font_size_hpc))
            rPr.append(sz)
            szCs = OxmlElement('w:szCs')
            szCs.set(qn('w:val'), str(font_size_hpc))
            rPr.append(szCs)
            if is_bold:
                rPr.append(OxmlElement('w:b'))
            
            # Make it look like a standard blue clickable link to make it OBVIOUS it's clickable
            col = OxmlElement('w:color')
            col.set(qn('w:val'), '0563C1') # Word's default hyperlink blue
            rPr.append(col)
            
            u = OxmlElement('w:u')
            u.set(qn('w:val'), 'single')
            rPr.append(u)
            
            return rPr
        
        # Hyperlink
        hl = OxmlElement('w:hyperlink')
        hl.set(qn('w:anchor'), bookmark_name)
        # Word tooltip instruction
        hl.set(qn('w:tooltip'), 'Ctrl+Click to follow link')
        
        hl_r = OxmlElement('w:r')
        hl_r.append(make_rPr())
        hl_t = OxmlElement('w:t')
        hl_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        hl_t.text = text
        hl_r.append(hl_t)
        hl.append(hl_r)
        
        p.append(hl)
        return p

    # Find insertion point (before Chapter I)
    chapter1_para = None
    for p in doc.paragraphs:
        if p.text.strip().upper().startswith('CHAPTER I'):
            chapter1_para = p
            break
            
    if not chapter1_para:
        print("Could not find CHAPTER I")
        return
        
    parent_body = chapter1_para._p.getparent()
    ch1_pos = list(parent_body).index(chapter1_para._p)
    
    toc_elements = []

    # TOC title
    toc_title_p = OxmlElement('w:p')
    toc_title_pPr = OxmlElement('w:pPr')
    toc_title_jc = OxmlElement('w:jc')
    toc_title_jc.set(qn('w:val'), 'center')
    toc_title_pPr.append(toc_title_jc)
    toc_title_spacing = OxmlElement('w:spacing')
    toc_title_spacing.set(qn('w:before'), '0')
    toc_title_spacing.set(qn('w:after'), '240')
    toc_title_pPr.append(toc_title_spacing)
    toc_title_p.append(toc_title_pPr)

    toc_title_rPr = OxmlElement('w:rPr')
    toc_title_rFonts = OxmlElement('w:rFonts')
    toc_title_rFonts.set(qn('w:ascii'), 'Times New Roman')
    toc_title_rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    toc_title_rPr.append(toc_title_rFonts)
    toc_title_sz = OxmlElement('w:sz')
    toc_title_sz.set(qn('w:val'), '32')
    toc_title_rPr.append(toc_title_sz)
    toc_title_rPr.append(OxmlElement('w:b'))
    toc_title_r = OxmlElement('w:r')
    toc_title_r.append(toc_title_rPr)
    toc_title_t = OxmlElement('w:t')
    toc_title_t.text = 'TABLE OF CONTENTS'
    toc_title_r.append(toc_title_t)
    toc_title_p.append(toc_title_r)
    toc_elements.append(toc_title_p)
    
    # Entries
    for (idx, text, level, bk_name) in toc_entries:
        toc_elements.append(create_toc_entry_para(text, level, bk_name))
        
    # Page break after TOC
    pb_p = OxmlElement('w:p')
    pb_r = OxmlElement('w:r')
    pb_br = OxmlElement('w:br')
    pb_br.set(qn('w:type'), 'page')
    pb_r.append(pb_br)
    pb_p.append(pb_r)
    toc_elements.append(pb_p)
    
    # Insert
    for offset, el in enumerate(toc_elements):
        parent_body.insert(ch1_pos + offset, el)
        
    print(f"Inserted manual TOC with {len(toc_entries)} clickable hyperlinks")

    doc.save(out_file)
    print(f"Saved to {out_file}")

rebuild_manual_toc(r'd:\RESEARCH PNGS\zukhruftheisis_cleaned.docx', r'd:\RESEARCH PNGS\zukhruftheisis_FINAL_TOC.docx')

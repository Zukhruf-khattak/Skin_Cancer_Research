"""
Changes to zukhruftheisis_revised.docx:
1. Page numbering: pre-intro = Roman numerals (i, ii, iii...), 
   from Introduction (Chapter I) = Arabic numerals starting at 1
2. Reduce references from 45 to 28 (keep refs 1-28, remove 29-45)
   and remove any in-text citations referencing removed refs [29]-[45]
3. Table/Figure caption text: set font to 10pt (Times New Roman)
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy, re

doc = Document(r'd:\RESEARCH PNGS\zukhruftheisis_revised.docx')

# =====================================================================
# HELPER: Apply font to a paragraph's runs
# =====================================================================
def apply_font(para, name='Times New Roman', size_pt=12, bold=None, color=(0,0,0)):
    for run in para.runs:
        run.font.name = name
        run.font.size = Pt(size_pt)
        if bold is not None:
            run.font.bold = bold
        if color is not None:
            run.font.color.rgb = RGBColor(*color)
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:ascii'), name)
        rFonts.set(qn('w:hAnsi'), name)
        rFonts.set(qn('w:cs'), name)

# =====================================================================
# STEP 1: FIX TABLE/FIGURE CAPTION FONT SIZE TO 10pt
# =====================================================================
print("Step 1: Setting Table/Figure caption font to 10pt...")
caption_count = 0
for para in doc.paragraphs:
    text = para.text.strip()
    # Match "Table X:" or "Figure X:" caption lines
    if re.match(r'^(Table|Figure)\s+[\d\.]+[\s:]', text, re.IGNORECASE):
        apply_font(para, name='Times New Roman', size_pt=10, bold=None, color=(0,0,0))
        caption_count += 1
        
print(f"  Updated {caption_count} table/figure captions to 10pt.")

# =====================================================================
# STEP 2: REDUCE REFERENCES FROM 45 TO 28
# References to REMOVE: [29] through [45]
# =====================================================================
print("\nStep 2: Reducing references to 28...")

# Find references section
refs_start_idx = -1
for i, para in enumerate(doc.paragraphs):
    if para.text.strip().upper() == 'REFERENCES':
        refs_start_idx = i
        break

print(f"  References section at paragraph index: {refs_start_idx}")

if refs_start_idx != -1:
    # Collect all reference paragraphs
    ref_paras = []
    for i in range(refs_start_idx + 1, len(doc.paragraphs)):
        para = doc.paragraphs[i]
        text = para.text.strip()
        if text:
            ref_paras.append((i, para, text))
    
    print(f"  Found {len(ref_paras)} reference entries")
    
    # Remove references [29] through [45] (indices 28-44 in the list, 0-based)
    # These are paragraphs where text starts with [29], [30], ..., [45]
    refs_to_remove = set(range(29, 46))  # refs 29-45
    
    to_delete = []
    for (idx, para, text) in ref_paras:
        m = re.match(r'^\[(\d+)\]', text)
        if m:
            ref_num = int(m.group(1))
            if ref_num in refs_to_remove:
                to_delete.append(para)
    
    print(f"  Removing {len(to_delete)} reference entries ([29] through [45])...")
    for para in to_delete:
        p = para._element
        p.getparent().remove(p)
    
    print("  References reduced to 28.")

# =====================================================================
# STEP 3: CLEAN UP IN-TEXT CITATIONS FOR REMOVED REFS [29]-[45]
# Replace [29]-[45] citations in body text with nothing or ellipsis
# =====================================================================
print("\nStep 3: Cleaning in-text citations for removed references...")

removed_refs_pattern = re.compile(r'\s*\[(?:' + '|'.join(str(n) for n in range(29, 46)) + r')\]')

cleaned_count = 0
for para in doc.paragraphs:
    for run in para.runs:
        original = run.text
        # Remove individual citations like [29], [30], etc.
        cleaned = re.sub(r'\s*\[(?:2[9]|3[0-9]|4[0-5])\]', '', original)
        if cleaned != original:
            run.text = cleaned
            cleaned_count += 1

print(f"  Cleaned {cleaned_count} in-text citation instances.")

# =====================================================================
# STEP 4: SET UP PAGE NUMBERING
# - Before Introduction (Chapter I): Roman numerals starting at i
# - From Introduction onwards: Arabic numerals starting at 1
#
# This requires inserting a section break (continuous or next page)
# just before the paragraph containing "CHAPTER I"
# =====================================================================
print("\nStep 4: Setting up page numbering...")

# Find "CHAPTER I" paragraph (the start of Introduction)
chapter1_idx = -1
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if re.match(r'^CHAPTER\s+I\b', text, re.IGNORECASE):
        chapter1_idx = i
        print(f"  Found 'CHAPTER I' at paragraph [{i}]: '{text}'")
        break

if chapter1_idx == -1:
    print("  WARNING: Could not find CHAPTER I")
else:
    # ---- Step 4a: Configure the body-level (final) sectPr for Roman numerals ----
    # This covers all pages BEFORE Chapter I
    body = doc.element.body
    body_sectPr = body.find(qn('w:sectPr'))
    if body_sectPr is None:
        body_sectPr = OxmlElement('w:sectPr')
        body.append(body_sectPr)
    
    # Set Roman numeral numbering starting at 1 for the front matter section
    # We'll do this by making the LAST section (body sectPr) Arabic from Chapter I,
    # and the sectPr we insert before Chapter I will have Roman numerals.
    
    # ---- Step 4b: Insert sectPr in the paragraph BEFORE Chapter I ----
    # We need to add a sectPr to the pPr of the paragraph JUST BEFORE Chapter I
    # This creates a section break, ending the front-matter section
    
    # The paragraph before Chapter I
    prev_para_idx = chapter1_idx - 1
    # Find the actual prev non-empty para
    while prev_para_idx > 0 and not doc.paragraphs[prev_para_idx].text.strip():
        prev_para_idx -= 1
    
    print(f"  Inserting section break after paragraph [{prev_para_idx}]: '{doc.paragraphs[prev_para_idx].text[:60]}'")
    
    prev_para = doc.paragraphs[prev_para_idx]
    prev_pPr = prev_para._p.find(qn('w:pPr'))
    if prev_pPr is None:
        prev_pPr = OxmlElement('w:pPr')
        prev_para._p.insert(0, prev_pPr)
    
    # Remove any existing sectPr in this pPr
    existing_sectPr = prev_pPr.find(qn('w:sectPr'))
    if existing_sectPr is not None:
        prev_pPr.remove(existing_sectPr)
    
    # Create new sectPr for the front-matter section (Roman numerals)
    front_sectPr = OxmlElement('w:sectPr')
    
    # Page numbering: Roman numerals, starting at 1
    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(qn('w:fmt'), 'lowerRoman')
    pgNumType.set(qn('w:start'), '1')
    front_sectPr.append(pgNumType)
    
    # Copy page size and margins from body sectPr so formatting is preserved
    for child_tag in [qn('w:pgSz'), qn('w:pgMar'), qn('w:cols')]:
        child = body_sectPr.find(child_tag)
        if child is not None:
            front_sectPr.append(copy.deepcopy(child))
    
    # Insert front_sectPr into prev paragraph's pPr
    prev_pPr.append(front_sectPr)
    
    # ---- Step 4c: Update body sectPr to Arabic numerals from 1 ----
    # This is the section starting at Chapter I
    pgNumType_main = body_sectPr.find(qn('w:pgNumType'))
    if pgNumType_main is not None:
        body_sectPr.remove(pgNumType_main)
    pgNumType_main = OxmlElement('w:pgNumType')
    pgNumType_main.set(qn('w:fmt'), 'decimal')
    pgNumType_main.set(qn('w:start'), '1')
    body_sectPr.append(pgNumType_main)
    
    print("  Section break and page numbering configured.")
    print("  Front matter: Roman numerals (i, ii, iii...)")
    print("  From Chapter I: Arabic numerals (1, 2, 3...)")

# =====================================================================
# ALSO: Add page numbers to headers/footers if not already present
# We need to ensure there are footer fields with PAGE number field
# =====================================================================
print("\nStep 5: Ensuring page number fields in footers...")

def add_page_number_footer(sectPr, num_format='decimal'):
    """Add a footer with centered page number to a sectPr."""
    # Check if footer already exists
    footerRef = sectPr.find(qn('w:footerReference'))
    if footerRef is not None:
        print(f"  Footer reference already exists in sectPr, skipping creation.")
        return
    
    # Create footer part - we need to add it via relationships
    # For simplicity, we'll check if the document already has footer parts
    print(f"  Note: Footer page number insertion requires manual Word setup or existing footer.")
    print(f"  The page numbering FORMAT is set correctly in sectPr.")
    print(f"  To see page numbers, add them via: Insert > Page Number in Word.")

# Check current footer refs in both sectPrs
body = doc.element.body
body_sectPr = body.find(qn('w:sectPr'))
print(f"  Body sectPr footer refs: {len(body_sectPr.findall(qn('w:footerReference')))}")
for i, para in enumerate(doc.paragraphs):
    pPr = para._p.find(qn('w:pPr'))
    if pPr is not None:
        s = pPr.find(qn('w:sectPr'))
        if s is not None:
            refs = s.findall(qn('w:footerReference'))
            print(f"  Para [{i}] sectPr footer refs: {len(refs)}")

# =====================================================================
# SAVE
# =====================================================================
output_path = r'd:\RESEARCH PNGS\zukhruftheisis_final.docx'
doc.save(output_path)
print(f"\nDONE! Saved to: {output_path}")
print("Summary of changes:")
print("  1. Table/Figure captions: 10pt Times New Roman")
print("  2. References: reduced from 45 to 28 (removed refs 29-45)")
print("  3. In-text citations [29]-[45]: removed from body text")
print("  4. Page numbering: front matter = Roman (i,ii,iii), Chapter I onwards = Arabic (1,2,3)")

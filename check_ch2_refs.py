import re
from docx import Document

INPUT_FILE = "Huma Thesis_Updated.docx"

try:
    doc = Document(INPUT_FILE)
except Exception as e:
    print(f"Error loading document: {e}")
    exit(1)

print("--- Chapter 2 Headings ---")
in_ch2 = False
ch2_paras = []
ref_section_start = -1

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    
    if text.upper() == "CHAPTER II":
        in_ch2 = True
    elif text.upper() == "CHAPTER III":
        in_ch2 = False
        
    if re.match(r'^(REFERENCES?|BIBLIOGRAPHY)', text, re.IGNORECASE) and len(text) < 30:
        ref_section_start = i
        
    if in_ch2:
        ch2_paras.append((i, text, para.style.name))
        # Print headings
        if re.match(r'^2\.\d+', text) and len(text) < 150:
            print(f"Para {i} | Style: {para.style.name} | Text: {text}")

print("\n--- Chapter 2 Citations ---")
citations = []
for i, text, style in ch2_paras:
    # Find citations like [1], [1, 2], [1-3]
    matches = re.findall(r'\[([0-9\s,\-]+)\]', text)
    for match in matches:
        # Check if it's purely numbers and separators
        if re.match(r'^[0-9\s,\-]+$', match):
            citations.append(match)

print(f"Found {len(citations)} citation blocks in Chapter 2. Some examples:")
print(citations[:20])

# Collect all unique numbers cited
cited_numbers = set()
for c in citations:
    parts = re.split(r'[,]', c)
    for p in parts:
        p = p.strip()
        if '-' in p:
            start, end = p.split('-')
            try:
                for num in range(int(start), int(end)+1):
                    cited_numbers.add(num)
            except ValueError:
                pass
        elif p.isdigit():
            cited_numbers.add(int(p))

if cited_numbers:
    print(f"Max cited number in Chapter 2: {max(cited_numbers)}")
    print(f"Total unique references cited in Chapter 2: {len(cited_numbers)}")

print("\n--- Bibliography Check ---")
if ref_section_start >= 0:
    ref_count = 0
    ref_numbers = []
    for i in range(ref_section_start + 1, len(doc.paragraphs)):
        text = doc.paragraphs[i].text.strip()
        match = re.match(r'^\[(\d+)\]', text)
        if match:
            ref_count += 1
            ref_numbers.append(int(match.group(1)))
    print(f"Found {ref_count} references in the bibliography.")
    if ref_numbers:
        print(f"Bibliography numbers range from {min(ref_numbers)} to {max(ref_numbers)}")
        # Check for missing citations in bibliography
        missing_in_bib = sorted(list(cited_numbers - set(ref_numbers)))
        if missing_in_bib:
            print(f"WARNING: The following numbers are cited in Chapter 2 but missing in Bibliography: {missing_in_bib}")
else:
    print("No References/Bibliography section found!")


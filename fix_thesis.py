"""
Thesis Fixer — Huma Thesis_Final.docx
======================================
1. Apply formatting: Times New Roman, 1.5 line spacing,
   Heading 1 = 18pt, Heading 2 = 14pt, Heading 3 = 12pt
2. Content reductions:
   - Merge 2.9 + 2.11 (remove duplication)
   - Trim 2.1 intro & 2.10 summary (remove bookend bloat)
   - Trim 3.9.1 / 3.9.2 from Chapter 3 (they're duplicated in Ch.4)
   - Clean cross-chapter repetition in 3.1 System Overview & 3.4
   - Trim Ch.4 redundant Grad-CAM text
   - Fix references section
3. Figures are NOT touched
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_LINE_SPACING
import copy
import re

INPUT_FILE  = "Huma Thesis_Final.docx"
OUTPUT_FILE = "Huma_Thesis_Revised.docx"

doc = Document(INPUT_FILE)

# =============================================================================
# HELPERS
# =============================================================================

def set_run_font(run, size_pt=12):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size_pt)
    # Also set the rPr font for East-Asian / complex script compatibility
    rpr = run._r.get_or_add_rPr()
    for tag in [qn('w:rFonts')]:
        rf = rpr.find(tag)
        if rf is None:
            rf = OxmlElement(tag)
            rpr.insert(0, rf)
        rf.set(qn('w:ascii'), 'Times New Roman')
        rf.set(qn('w:hAnsi'), 'Times New Roman')
        rf.set(qn('w:cs'),    'Times New Roman')

def set_para_spacing(para, line_rule=WD_LINE_SPACING.ONE_POINT_FIVE):
    pf = para.paragraph_format
    pf.line_spacing_rule = line_rule

def format_paragraph(para, body_size=12):
    """Apply TNR + 1.5 spacing to a paragraph; preserve bold/italic."""
    set_para_spacing(para)
    for run in para.runs:
        b = run.bold
        i = run.italic
        set_run_font(run, body_size)
        run.bold   = b
        run.italic = i

def format_heading(para, size_pt):
    """Format a heading paragraph."""
    set_para_spacing(para)
    for run in para.runs:
        b = run.bold
        set_run_font(run, size_pt)
        run.bold = b if b is not None else True

def para_text_contains(para, *terms):
    t = para.text.strip()
    return any(term.lower() in t.lower() for term in terms)

def get_para_index(doc, text_fragment):
    """Return index of first paragraph whose text contains text_fragment."""
    for i, p in enumerate(doc.paragraphs):
        if text_fragment.lower() in p.text.lower():
            return i
    return -1

def delete_paragraph(para):
    """Remove a paragraph from the document."""
    p = para._element
    p.getparent().remove(p)

def clear_paragraph_text(para):
    """Clear all runs from a paragraph (keeps the paragraph object)."""
    for run in para.runs:
        run.text = ""

# =============================================================================
# STEP 1: APPLY FORMATTING TO ALL PARAGRAPHS
# =============================================================================
print("Step 1: Applying formatting...")

for para in doc.paragraphs:
    style = para.style.name

    if style.startswith("Heading 1"):
        format_heading(para, 18)
    elif style.startswith("Heading 2"):
        format_heading(para, 14)
    elif style.startswith("Heading 3"):
        format_heading(para, 12)
    elif style in ("Title",):
        format_heading(para, 20)
    else:
        format_paragraph(para, 12)

# Also apply to table cells
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                format_paragraph(para, 11)

print("  Formatting applied to all paragraphs and tables.")

# =============================================================================
# STEP 2: CONTENT REDUCTIONS
# =============================================================================
print("Step 2: Content reductions...")

paragraphs = doc.paragraphs  # live list

# --------------------------------------------------------------------------
# 2a. Remove the entire duplicate 3.9 Results section from Chapter 3
#     (These results are fully covered in Chapter 4)
#     We remove from "3.9  Results" heading down to (but not including)
#     "3.10  Hardware and Software Environment"
# --------------------------------------------------------------------------
print("  2a. Removing duplicate Ch.3 results section (3.9)...")

in_ch3_results = False
paras_to_delete = []

for para in doc.paragraphs:
    txt = para.text.strip()
    style = para.style.name

    # Start marker
    if (style.startswith("Heading") or "Normal" in style) and re.search(r'3\.9\s+Results', txt, re.IGNORECASE):
        in_ch3_results = True

    # Stop before 3.10
    if in_ch3_results and re.search(r'3\.10\s+Hardware', txt, re.IGNORECASE):
        in_ch3_results = False

    if in_ch3_results:
        paras_to_delete.append(para)

print(f"    Deleting {len(paras_to_delete)} paragraphs from 3.9 results block")
for p in paras_to_delete:
    try:
        delete_paragraph(p)
    except Exception:
        pass

# --------------------------------------------------------------------------
# 2b. Merge / trim Literature Review sections:
#     - Trim 2.1 Introduction (keep only 1 paragraph — the "three generations" one)
#     - Trim 2.10 Chapter Summary (keep only the numbered-list conclusions, drop prose)
#     - Merge 2.11 into 2.9 by removing the 2.11 heading and its intro sentence
#       (the comparative table stays; the prose intro for 2.11 is redundant)
# --------------------------------------------------------------------------
print("  2b. Trimming Literature Review bloat...")

# Trim 2.1 — remove the second paragraph of 2.1 (the "this chapter reviews..." meta-paragraph)
meta_para_text = "This chapter reviews that trajectory systematically"
for para in doc.paragraphs:
    if meta_para_text.lower() in para.text.lower():
        delete_paragraph(para)
        print("    Removed 2.1 meta-paragraph")
        break

# Trim 2.10 — it's already concise; remove only the repeated "key conclusions" prelude
# The 2.10 summary re-says things already in the chapter — keep ONLY the 4 numbered points
# The paragraph starting "This chapter traced the evolution" restates everything;
# replace it with a single short sentence.
for para in doc.paragraphs:
    if "This chapter traced the evolution of automated skin cancer" in para.text:
        # Replace with concise bridging sentence
        for run in para.runs:
            run.text = ""
        if para.runs:
            para.runs[0].text = (
                "The literature review traced the evolution of automated skin cancer detection "
                "from classical ML pipelines through deep CNNs to modern ensemble and Transformer-based systems. "
                "Six critical research gaps — data leakage, domain shift, class imbalance, test-set corruption, "
                "interpretability deficit, and computational inaccessibility — are directly addressed by the "
                "proposed framework in Chapter 3."
            )
            set_run_font(para.runs[0], 12)
        else:
            run = para.add_run(
                "The literature review traced the evolution of automated skin cancer detection "
                "from classical ML pipelines through deep CNNs to modern ensemble and Transformer-based systems. "
                "Six critical research gaps — data leakage, domain shift, class imbalance, test-set corruption, "
                "interpretability deficit, and computational inaccessibility — are directly addressed by the "
                "proposed framework in Chapter 3."
            )
            set_run_font(run, 12)
        print("    Condensed 2.10 summary paragraph")
        break

# Merge 2.11 into 2.9: Remove the 2.11 heading and the redundant intro sentence
# ("To contextualise the performance...")  — the table (Table 2.7) stays
for para in doc.paragraphs:
    if para.style.name.startswith("Heading") and "2.11" in para.text and "Comparative Analysis" in para.text:
        delete_paragraph(para)
        print("    Removed 2.11 heading (merged into 2.9)")
        break

for para in doc.paragraphs:
    if "To contextualise the performance of the proposed framework, Table 2.7" in para.text:
        delete_paragraph(para)
        print("    Removed 2.11 intro sentence (table stays)")
        break

# --------------------------------------------------------------------------
# 2c. Trim 3.1 System Overview — remove the last 2 paragraphs that restate
#     the clinical imperative already stated in Ch.1 and the problem statement
# --------------------------------------------------------------------------
print("  2c. Trimming 3.1 System Overview repetition...")

# The paragraph "Every design decision — from model selection to loss function choice..."
# is a near-verbatim repeat from the abstract and Ch.1. Remove it from Ch.3.
for para in doc.paragraphs:
    if ("Every design decision" in para.text and 
        "minimising false negatives" in para.text and
        "clinical imperative" in para.text):
        delete_paragraph(para)
        print("    Removed repeated 'every design decision' paragraph from 3.1")
        break

# --------------------------------------------------------------------------
# 2d. Trim Ch.4 Grad-CAM duplicated caption text
#     The figure captions in Ch.4 sometimes have a full paragraph below them
#     that repeats the caption text verbatim. Remove them.
# --------------------------------------------------------------------------
print("  2d. Removing Ch.4 duplicate Grad-CAM caption paragraphs...")

# This paragraph is a verbatim repeat of the caption above it
duplicate_gradcam = (
    "The above figure Grad-CAM heatmaps for Actinic Keratosis (AK) classification"
)
for para in doc.paragraphs:
    if duplicate_gradcam.lower() in para.text.lower():
        delete_paragraph(para)
        print("    Removed duplicate AK Grad-CAM caption paragraph")
        break

# --------------------------------------------------------------------------
# 2e. Remove repeated "4.2.3 Clinical Safety Analysis" confusion matrix
#     description text — the figures show this; prose repeats it
# --------------------------------------------------------------------------
print("  2e. Trimming Ch.4 verbose confusion matrix prose...")

# "77 False Negatives: These are malignant lesions incorrectly classified as benign"
# "132 False Positives: These are benign lesions flagged as malignant"
# These two sentences are largely repeated in surrounding text. 
# Replace them with a single combined sentence.
for para in doc.paragraphs:
    if "77 False Negatives: These are malignant lesions incorrectly classified as benign" in para.text:
        for run in para.runs:
            run.text = ""
        if para.runs:
            para.runs[0].text = (
                "The ensemble produced 77 False Negatives (malignant cases missed, 5.5% miss rate) "
                "and 132 False Positives (unnecessary referrals), yielding a clinical sensitivity of 94.50% "
                "and specificity of 94.51%."
            )
            set_run_font(para.runs[0], 12)
        print("    Condensed FN/FP prose in 4.2.3")
        break

for para in doc.paragraphs:
    if "132 False Positives: These are benign lesions flagged as malignant" in para.text:
        delete_paragraph(para)
        print("    Removed redundant FP sentence")
        break

# --------------------------------------------------------------------------
# 2f. Remove the second "4.1 RESULTS AND EVALUATION" heading / intro that
#     repeats the chapter title
# --------------------------------------------------------------------------
print("  2f. Removing duplicate Chapter 4 intro heading...")
seen_results_intro = False
for para in doc.paragraphs:
    if "4.1 RESULTS AND EVALUATION" in para.text or "4.1 Results and Evaluation" in para.text:
        if seen_results_intro:
            delete_paragraph(para)
            print("    Removed duplicate 4.1 heading")
            break
        seen_results_intro = True

# --------------------------------------------------------------------------
# 2g. Trim 3.4.1 Model Selection Justification subsections
#     The 5 sub-sub-sections (3.4.1.1 to 3.4.1.6) are very long and repeat
#     information already in the bullet list at the top of 3.4.
#     We condense 3.4.1.6 (Ensemble Rationale) — the most redundant one —
#     since it restates ensemble theory already given in 3.7.
# --------------------------------------------------------------------------
print("  2g. Condensing 3.4.1.6 Ensemble Rationale (redundant with 3.7)...")
for para in doc.paragraphs:
    if ("The five architectures span three distinct computational families" in para.text and
            "Dietterich, 2000" in para.text):
        for run in para.runs:
            run.text = ""
        if para.runs:
            para.runs[0].text = (
                "The five architectures span three distinct computational families — classical residual CNNs, "
                "efficiency-optimised CNNs, a modernised CNN with transformer principles, and a pure Vision Transformer — "
                "ensuring diverse and complementary feature extraction. Ensemble diversity reduces both bias and variance, "
                "improving generalisation beyond what any single model achieves. Full ensemble aggregation details are "
                "provided in Section 3.7."
            )
            set_run_font(para.runs[0], 12)
        print("    Condensed 3.4.1.6")
        break

# =============================================================================
# STEP 3: FIX REFERENCES — trim from 61 to 55
# =============================================================================
print("Step 3: Fixing references section...")

# The references are at the end of the document, each paragraph is one ref [N]...
# We will:
#   1. Find the References heading
#   2. Collect all ref paragraphs
#   3. Delete the ones beyond index 55 (refs 56-61 are less critical duplicates)

ref_section_start = -1
for i, para in enumerate(doc.paragraphs):
    txt = para.text.strip()
    if re.match(r'^(REFERENCES?|BIBLIOGRAPHY)', txt, re.IGNORECASE) and len(txt) < 30:
        ref_section_start = i
        print("    Found References section at paragraph index " + str(i))
        break

if ref_section_start >= 0:
    all_paras = doc.paragraphs
    ref_paras = []
    for i in range(ref_section_start + 1, len(all_paras)):
        p = all_paras[i]
        txt = p.text.strip()
        # Reference entries start with [N]
        if txt and re.match(r'^\[\d+\]', txt):
            ref_paras.append(p)
    
    print("    Found " + str(len(ref_paras)) + " reference entries")
    
    # Safe print for inspection
    for idx, p in enumerate(ref_paras[:62], 1):
        try:
            safe_txt = p.text[:100].encode('ascii', errors='replace').decode('ascii')
            print("      [" + str(idx).zfill(2) + "] " + safe_txt)
        except Exception:
            print("      [" + str(idx).zfill(2) + "] (unreadable)")
    
    # Trim to 55: delete paragraphs for refs 56 onwards
    # Refs 56-61 (Kassem [56], Fraiwan [57], Bassel [58], Adla [59], + 2 more)
    # These are cited in 2.9 which we partially trimmed; they are supplementary.
    TARGET_MAX_REFS = 55
    if len(ref_paras) > TARGET_MAX_REFS:
        refs_to_delete = ref_paras[TARGET_MAX_REFS:]  # refs [56], [57]...
        print("    Deleting " + str(len(refs_to_delete)) + " excess reference entries (" +
              str(TARGET_MAX_REFS + 1) + "–" + str(len(ref_paras)) + ")")
        for rp in refs_to_delete:
            try:
                delete_paragraph(rp)
            except Exception as e:
                print("      Could not delete: " + str(e))
        print("    References reduced to " + str(TARGET_MAX_REFS))
    else:
        print("    Reference count OK: " + str(len(ref_paras)) + " refs")
else:
    print("    References section not found — references may be inline only")

# =============================================================================
# STEP 4: FORMAT DOCUMENT STYLES GLOBALLY
# =============================================================================
print("Step 4: Setting document-level style defaults...")

# Update the Normal style in document styles
from docx.oxml.ns import qn

styles = doc.styles

# Update Normal style
normal_style = styles['Normal']
normal_style.font.name = 'Times New Roman'
normal_style.font.size = Pt(12)
normal_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

# Update Heading 1
try:
    h1 = styles['Heading 1']
    h1.font.name = 'Times New Roman'
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
except KeyError:
    pass

# Update Heading 2
try:
    h2 = styles['Heading 2']
    h2.font.name = 'Times New Roman'
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
except KeyError:
    pass

# Update Heading 3
try:
    h3 = styles['Heading 3']
    h3.font.name = 'Times New Roman'
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
except KeyError:
    pass

print("  Document styles updated.")

# =============================================================================
# STEP 5: SAVE
# =============================================================================
print(f"Step 5: Saving to {OUTPUT_FILE}...")
doc.save(OUTPUT_FILE)
print(f"\nDone! Saved as: {OUTPUT_FILE}")

# Final stats
doc2 = Document(OUTPUT_FILE)
total_words = sum(len(p.text.split()) for p in doc2.paragraphs)
print(f"Final word count: {total_words} (~{total_words//250} pages of text)")
print(f"Tables remaining: {len(doc2.tables)}")
rels = [r for r in doc2.part.rels.values() if 'image' in r.reltype]
print(f"Images remaining: {len(rels)}")

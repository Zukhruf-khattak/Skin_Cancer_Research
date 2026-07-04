import docx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def rebuild_clean_toc(in_file, out_file):
    doc = docx.Document(in_file)
    
    # 1. Remove old TOC paragraphs
    paras_to_remove = []
    for p in doc.paragraphs:
        text = p.text.strip().upper()
        if 'TABLE OF CONTENTS' in text:
            paras_to_remove.append(p)
            
        hls = p._p.findall('.//' + qn('w:hyperlink'))
        if hls and any(hl.get(qn('w:anchor'), '').startswith('TOC_BM_') for hl in hls):
            paras_to_remove.append(p)
            
    for p in paras_to_remove:
        p._element.getparent().remove(p._element)

    # 2. Collect headings and ensure bookmarks
    toc_entries = []
    bk_counter = 0
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        style = para.style.name
        if not text:
            continue
        
        level = None
        if re.match(r'^CHAPTER\s+(I|II|III|IV|V|VI|VII|VIII|IX|X)\b', text, re.IGNORECASE):
            level = 1
        elif text.upper() == 'REFERENCES':
            level = 1
        elif 'Heading 2' in style:
            level = 2
        elif 'Heading 3' in style:
            level = 3
        
        if level is not None:
            bk_name = f'TOC_BM_{bk_counter}'
            bk_counter += 1
            toc_entries.append((i, text, level, bk_name))

    # Add bookmarks to heading paragraphs
    for (idx, text, level, bk_name) in toc_entries:
        para = doc.paragraphs[idx]
        for bk in para._p.findall('.//' + qn('w:bookmarkStart')):
            if bk.get(qn('w:name'), '').startswith('TOC_BM_'):
                bk.getparent().remove(bk)
        for bk in para._p.findall('.//' + qn('w:bookmarkEnd')):
            bk.getparent().remove(bk)
        
        bk_id = str(9000 + idx)
        bk_start = OxmlElement('w:bookmarkStart')
        bk_start.set(qn('w:id'), bk_id)
        bk_start.set(qn('w:name'), bk_name)
        bk_end = OxmlElement('w:bookmarkEnd')
        bk_end.set(qn('w:id'), bk_id)
        
        para._p.insert(0, bk_start)
        para._p.append(bk_end)

    # 3. Build valid TOC elements
    def create_toc_entry_para(text, level, bookmark_name):
        p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        
        indent_twips = {1: 0, 2: 360, 3: 720}.get(level, 0)
        if indent_twips:
            ind = OxmlElement('w:ind')
            ind.set(qn('w:left'), str(indent_twips))
            pPr.append(ind)
        
        # Tab stop for page number
        tabs = OxmlElement('w:tabs')
        tab = OxmlElement('w:tab')
        tab.set(qn('w:val'), 'right')
        tab.set(qn('w:leader'), 'dot')
        tab.set(qn('w:pos'), '9000') # ~6.25 inches
        tabs.append(tab)
        pPr.append(tabs)
        
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:before'), '40')
        spacing.set(qn('w:after'), '40')
        pPr.append(spacing)
        
        p.append(pPr)
        
        font_size_hpc = {1: 24, 2: 24, 3: 24}.get(level, 24)
        is_bold = (level == 1)
        
        def make_rPr():
            rPr = OxmlElement('w:rPr')
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), 'Times New Roman')
            rFonts.set(qn('w:hAnsi'), 'Times New Roman')
            rPr.append(rFonts)
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), str(font_size_hpc))
            rPr.append(sz)
            szCs = OxmlElement('w:szCs')
            szCs.set(qn('w:val'), str(font_size_hpc))
            rPr.append(szCs)
            if is_bold:
                rPr.append(OxmlElement('w:b'))
            
            # BLACK TEXT
            col = OxmlElement('w:color')
            col.set(qn('w:val'), '000000')
            rPr.append(col)
            return rPr
        
        # --- A. HYPERLINK (Text only) ---
        hl = OxmlElement('w:hyperlink')
        hl.set(qn('w:anchor'), bookmark_name)
        hl.set(qn('w:tooltip'), 'Ctrl+Click to follow link')
        
        hl_r = OxmlElement('w:r')
        hl_r.append(make_rPr())
        hl_t = OxmlElement('w:t')
        hl_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        hl_t.text = text
        hl_r.append(hl_t)
        hl.append(hl_r)
        
        p.append(hl)  # Add hyperlink to paragraph
        
        # --- B. PAGE NUMBER (Outside hyperlink) ---
        # Add the tab character
        tab_r = OxmlElement('w:r')
        tab_r.append(make_rPr())
        tab_el = OxmlElement('w:tab')
        tab_r.append(tab_el)
        p.append(tab_r)
        
        # Add PAGEREF field
        r_fld_begin = OxmlElement('w:r')
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        r_fld_begin.append(fldChar1)
        p.append(r_fld_begin)
        
        r_instr = OxmlElement('w:r')
        r_instr.append(make_rPr())
        instr = OxmlElement('w:instrText')
        instr.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        instr.text = f' PAGEREF {bookmark_name} \\h '
        r_instr.append(instr)
        p.append(r_instr)
        
        r_fld_sep = OxmlElement('w:r')
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        r_fld_sep.append(fldChar2)
        p.append(r_fld_sep)
        
        # Default text until updated
        r_def = OxmlElement('w:r')
        r_def.append(make_rPr())
        t_def = OxmlElement('w:t')
        t_def.text = '1'
        r_def.append(t_def)
        p.append(r_def)
        
        r_fld_end = OxmlElement('w:r')
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        r_fld_end.append(fldChar3)
        p.append(r_fld_end)
        
        return p

    # Find insertion point
    chapter1_para = None
    for p in doc.paragraphs:
        if p.text.strip().upper().startswith('CHAPTER I'):
            chapter1_para = p
            break
            
    if not chapter1_para:
        print("Could not find CHAPTER I")
        return
        
    parent_body = chapter1_para._p.getparent()
    ch1_pos = list(parent_body).index(chapter1_para._p)
    
    toc_elements = []

    # TOC title
    toc_title_p = OxmlElement('w:p')
    toc_title_pPr = OxmlElement('w:pPr')
    toc_title_jc = OxmlElement('w:jc')
    toc_title_jc.set(qn('w:val'), 'center')
    toc_title_pPr.append(toc_title_jc)
    toc_title_spacing = OxmlElement('w:spacing')
    toc_title_spacing.set(qn('w:before'), '0')
    toc_title_spacing.set(qn('w:after'), '240')
    toc_title_pPr.append(toc_title_spacing)
    toc_title_p.append(toc_title_pPr)

    toc_title_rPr = OxmlElement('w:rPr')
    toc_title_rFonts = OxmlElement('w:rFonts')
    toc_title_rFonts.set(qn('w:ascii'), 'Times New Roman')
    toc_title_rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    toc_title_rPr.append(toc_title_rFonts)
    toc_title_sz = OxmlElement('w:sz')
    toc_title_sz.set(qn('w:val'), '32')
    toc_title_rPr.append(toc_title_sz)
    toc_title_rPr.append(OxmlElement('w:b'))
    toc_title_r = OxmlElement('w:r')
    toc_title_r.append(toc_title_rPr)
    toc_title_t = OxmlElement('w:t')
    toc_title_t.text = 'TABLE OF CONTENTS'
    toc_title_r.append(toc_title_t)
    toc_title_p.append(toc_title_r)
    toc_elements.append(toc_title_p)
    
    # Entries
    for (idx, text, level, bk_name) in toc_entries:
        toc_elements.append(create_toc_entry_para(text, level, bk_name))
        
    # Page break after TOC
    pb_p = OxmlElement('w:p')
    pb_r = OxmlElement('w:r')
    pb_br = OxmlElement('w:br')
    pb_br.set(qn('w:type'), 'page')
    pb_r.append(pb_br)
    pb_p.append(pb_r)
    toc_elements.append(pb_p)
    
    # Insert
    for offset, el in enumerate(toc_elements):
        parent_body.insert(ch1_pos + offset, el)
        
    # Ensure updateFields is set
    try:
        settings = doc.settings._element
        updateFields = settings.find(qn('w:updateFields'))
        if updateFields is None:
            updateFields = OxmlElement('w:updateFields')
            updateFields.set(qn('w:val'), 'true')
            settings.append(updateFields)
    except Exception as e:
        print(f"Settings error: {e}")

    doc.save(out_file)
    print(f"Saved to {out_file}")

rebuild_clean_toc(r'd:\RESEARCH PNGS\zukhruftheisis_Black_TOC.docx', r'd:\RESEARCH PNGS\zukhruftheisis_Final_Working_TOC.docx')

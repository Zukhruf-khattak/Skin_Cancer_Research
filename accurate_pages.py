import docx
from docx.oxml.ns import qn

def accurate_page_numbers(filename):
    doc = docx.Document(filename)
    
    current_page = 1
    page_format = 'roman' # start with roman
    
    # First find all the bookmark paragraphs
    bkmk_to_para = {}
    for i, p in enumerate(doc.paragraphs):
        for bk in p._p.findall('.//' + qn('w:bookmarkStart')):
            name = bk.get(qn('w:name'), '')
            if name.startswith('TOC_BM_'):
                bkmk_to_para[name] = (i, p.text.strip())
                
    # Now simulate pagination
    # We will look at sections to see where page numbers reset
    para_to_sect = []
    current_sect = 0
    for i, p in enumerate(doc.paragraphs):
        pPr = p._p.find(qn('w:pPr'))
        sectPr = None
        if pPr is not None:
            sectPr = pPr.find(qn('w:sectPr'))
        if sectPr is not None:
            para_to_sect.append(current_sect)
            current_sect += 1
        else:
            para_to_sect.append(current_sect)
            
    # Print section info
    for si, section in enumerate(doc.sections):
        sectPr = section._sectPr
        pnt = sectPr.find(qn('w:pgNumType'))
        fmt = pnt.get(qn('w:fmt')) if pnt is not None else 'none'
        start = pnt.get(qn('w:start')) if pnt is not None else 'none'
        print(f"Section {si}: fmt={fmt}, start={start}")
        
    print("\nScanning paragraphs for rendered page breaks...")
    page_display = 1
    sect_idx = 0
    
    mapping = {}
    
    for i, p in enumerate(doc.paragraphs):
        # Update section if we moved to a new one
        my_sect = para_to_sect[i]
        if my_sect > sect_idx:
            # We entered a new section. Check if it resets page number
            sect_idx = my_sect
            if sect_idx < len(doc.sections):
                sPr = doc.sections[sect_idx]._sectPr
                pnt = sPr.find(qn('w:pgNumType'))
                start = pnt.get(qn('w:start')) if pnt is not None else None
                if start is not None:
                    try:
                        page_display = int(start)
                        print(f"  [Para {i}] Section {sect_idx} reset page to {page_display}")
                    except:
                        pass
        
        # Check if this paragraph is a heading we want
        for bk in p._p.findall('.//' + qn('w:bookmarkStart')):
            name = bk.get(qn('w:name'), '')
            if name.startswith('TOC_BM_'):
                mapping[name] = page_display
                if 'TOC_BM_0' in name or 'TOC_BM_1' in name:
                    print(f"  --> Heading {name} '{p.text.strip()[:30]}' is on page {page_display}")
                    
        # Check for breaks IN this paragraph to increment FOR THE NEXT paragraph
        breaks = p._p.findall('.//' + qn('w:br'))
        rendered = p._p.findall('.//' + qn('w:lastRenderedPageBreak'))
        
        num_breaks = sum(1 for br in breaks if br.get(qn('w:type')) == 'page')
        num_breaks += len(rendered)
        
        # If this paragraph HAS a section break, does it also act as a page break?
        # Usually yes, SectionType.NEW_PAGE.
        pPr = p._p.find(qn('w:pPr'))
        if pPr is not None and pPr.find(qn('w:sectPr')) is not None:
            # Check section type of the NEXT section
            if sect_idx + 1 < len(doc.sections):
                nxt_sPr = doc.sections[sect_idx + 1]._sectPr
                type_el = nxt_sPr.find(qn('w:type'))
                if type_el is None or type_el.get(qn('w:val')) != 'continuous':
                    num_breaks += 1
                    
        page_display += num_breaks
        
    print("\nSample mapped headings:")
    for i in range(5):
        name = f"TOC_BM_{i}"
        if name in mapping:
            print(f"{name}: Page {mapping[name]}")

accurate_page_numbers(r'd:\RESEARCH PNGS\zukhruftheisis_FINAL_TOC.docx')

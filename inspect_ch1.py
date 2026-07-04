import docx
from docx.oxml.ns import qn

def inspect_chapter_1(filename):
    doc = docx.Document(filename)
    ch1_idx = -1
    for i, p in enumerate(doc.paragraphs):
        if 'CHAPTER I' in p.text.upper():
            ch1_idx = i
            break
            
    print(f"Chapter 1 is at para {ch1_idx}")
    
    for i in range(ch1_idx - 5, ch1_idx + 10):
        p = doc.paragraphs[i]
        text = p.text.strip()
        breaks = p._p.findall('.//' + qn('w:br'))
        rendered = p._p.findall('.//' + qn('w:lastRenderedPageBreak'))
        sect = False
        pPr = p._p.find(qn('w:pPr'))
        if pPr is not None and pPr.find(qn('w:sectPr')) is not None:
            sect = True
            
        b_types = [br.get(qn('w:type'), 'textWrapping') for br in breaks]
        print(f"Para {i}: txt='{text[:30]}' | br={b_types} | rendered={len(rendered)} | sect={sect}")

inspect_chapter_1(r'd:\RESEARCH PNGS\zukhruftheisis_FINAL_TOC.docx')

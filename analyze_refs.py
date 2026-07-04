from docx import Document
from docx.oxml.ns import qn
import re
from lxml import etree

doc = Document(r'd:\RESEARCH PNGS\zukhruftheisis_final.docx')

# Get full XML text of every paragraph (including field codes, hyperlinks etc.)
print("=== PARAGRAPHS WITH [N] CITATIONS (XML scan) ===\n")

refs_start = -1
for i, p in enumerate(doc.paragraphs):
    if p.text.strip().upper() == 'REFERENCES':
        refs_start = i
        break

citations_all = {}  # ref_num -> [(para_idx, context)]

for i in range(refs_start):
    para = doc.paragraphs[i]
    # Get ALL text including from hyperlinks and field runs
    full_text = ''.join(t.text or '' for t in para._p.findall('.//' + qn('w:t')))
    
    matches = re.findall(r'\[(\d+)\]', full_text)
    if matches:
        for m in matches:
            num = int(m)
            if num not in citations_all:
                citations_all[num] = []
            citations_all[num].append((i, full_text.strip()[:100]))
        
print("All citations found:", sorted(citations_all.keys()))
print("Total unique:", len(citations_all))

print()
for num in sorted(citations_all.keys()):
    for (idx, ctx) in citations_all[num][:2]:
        print(f"  [{num}] at para {idx}: {ctx[:100]}")

# Also check what's in the references
print("\n\n=== ALL REFS IN BIBLIOGRAPHY ===")
refs = {}
for i in range(refs_start+1, len(doc.paragraphs)):
    full_text = ''.join(t.text or '' for t in doc.paragraphs[i]._p.findall('.//' + qn('w:t')))
    full_text = full_text.strip()
    if full_text:
        m = re.match(r'^\[(\d+)\](.+)', full_text)
        if m:
            refs[int(m.group(1))] = m.group(2).strip()

print(f"References: {sorted(refs.keys())}")

# Cross-check
cited_set = set(citations_all.keys())
ref_set = set(refs.keys())
print(f"\nCited but NOT in bibliography: {sorted(cited_set - ref_set)}")
print(f"In bibliography but NOT cited: {sorted(ref_set - cited_set)}")
print(f"Properly matched: {sorted(cited_set & ref_set)}")

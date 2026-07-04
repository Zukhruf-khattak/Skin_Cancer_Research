import docx
from docx.oxml.ns import qn

def detail_paras(filename):
    doc = docx.Document(filename)
    print("=== Paras 40 to 60 ===")
    for i in range(40, min(60, len(doc.paragraphs))):
        p = doc.paragraphs[i]
        text = p.text.strip()
        
        # Check for page breaks
        breaks = p._p.findall('.//' + qn('w:br'))
        has_pb = any(br.get(qn('w:type')) == 'page' for br in breaks)
        
        # Check for section breaks in this paragraph
        sect_break = False
        pPr = p._p.find(qn('w:pPr'))
        if pPr is not None and pPr.find(qn('w:sectPr')) is not None:
            sect_break = True
            
        print(f"[{i}] len={len(text)} PB={has_pb} SectBrk={sect_break} | {text[:50]}")

detail_paras(r'd:\RESEARCH PNGS\zukhruftheisis_final_version.docx')

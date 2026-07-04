import docx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

def add_native_toc(doc_path, output_path):
    doc = docx.Document(doc_path)
    
    # 1. Remove my manual TOC
    # The manual TOC started with a paragraph containing "TABLE OF CONTENTS"
    # and ended right before "CHAPTER I"
    toc_start_idx = -1
    chapter1_idx = -1
    
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip().upper()
        if 'TABLE OF CONTENTS' in text and toc_start_idx == -1:
            toc_start_idx = i
        if 'CHAPTER I' in text and chapter1_idx == -1:
            chapter1_idx = i
            
    print(f"TOC start: {toc_start_idx}, Chapter I: {chapter1_idx}")
    
    if toc_start_idx != -1 and chapter1_idx != -1:
        # Delete paragraphs between toc_start_idx and chapter1_idx - 1
        # Wait, the page break was inserted in the paragraph before TOC, or after?
        # Let's delete all paragraphs from toc_start_idx to chapter1_idx - 1
        
        # Actually, let's just find any paragraph with a hyperlink to "_Toc" and remove it
        paras_to_remove = []
        for p in doc.paragraphs:
            # Check for hyperlink with anchor starting with _Toc
            hls = p._p.findall('.//' + qn('w:hyperlink'))
            if hls and any(hl.get(qn('w:anchor'), '').startswith('_Toc_') for hl in hls):
                paras_to_remove.append(p)
            elif p.text.strip().upper() == 'TABLE OF CONTENTS':
                paras_to_remove.append(p)
                
        print(f"Removing {len(paras_to_remove)} manual TOC paragraphs")
        for p in paras_to_remove:
            p._element.getparent().remove(p._element)

    # 2. Add Native TOC field right before Chapter I
    # Find Chapter I again after removal
    chapter1_para = None
    for p in doc.paragraphs:
        if p.text.strip().upper().startswith('CHAPTER I'):
            chapter1_para = p
            break
            
    if chapter1_para is not None:
        parent = chapter1_para._p.getparent()
        ch1_index = list(parent).index(chapter1_para._p)
        
        # We will insert a native TOC field here.
        # It consists of:
        # Paragraph 1: "TABLE OF CONTENTS" heading
        # Paragraph 2: TOC field code
        
        # Heading
        p_head = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        pStyle = OxmlElement('w:pStyle')
        pStyle.set(qn('w:val'), 'Heading1') # or something appropriate
        pPr.append(pStyle)
        
        # Just manually style it to be 16pt bold center
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'center')
        pPr.append(jc)
        
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:before'), '0')
        spacing.set(qn('w:after'), '240')
        pPr.append(spacing)
        p_head.append(pPr)
        
        rPr = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        rPr.append(rFonts)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '32')
        rPr.append(sz)
        b = OxmlElement('w:b')
        rPr.append(b)
        
        r = OxmlElement('w:r')
        r.append(rPr)
        t = OxmlElement('w:t')
        t.text = 'TABLE OF CONTENTS'
        r.append(t)
        p_head.append(r)
        
        parent.insert(ch1_index, p_head)
        
        # TOC Field Code paragraph
        p_toc = OxmlElement('w:p')
        
        # Begin field
        r1 = OxmlElement('w:r')
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        r1.append(fldChar1)
        p_toc.append(r1)
        
        # InstrText
        r2 = OxmlElement('w:r')
        instr = OxmlElement('w:instrText')
        instr.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        # \o "1-3" : outlines level 1 to 3
        # \h : hyperlinks
        # \z : hide tab leader and page numbers in web layout
        # \u : use applied paragraph outline level
        instr.text = ' TOC \\o "1-3" \\h \\z \\u '
        r2.append(instr)
        p_toc.append(r2)
        
        # Separate field
        r3 = OxmlElement('w:r')
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        r3.append(fldChar2)
        p_toc.append(r3)
        
        # End field
        r4 = OxmlElement('w:r')
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        r4.append(fldChar3)
        p_toc.append(r4)
        
        parent.insert(ch1_index + 1, p_toc)
        
        # Page break after TOC
        p_break = OxmlElement('w:p')
        r_break = OxmlElement('w:r')
        br = OxmlElement('w:br')
        br.set(qn('w:type'), 'page')
        r_break.append(br)
        p_break.append(r_break)
        parent.insert(ch1_index + 2, p_break)
        
        print("Inserted Native TOC field.")

    # 3. Modify settings.xml to update fields on open
    try:
        from docx.oxml.core import oxml_tostring
        settings = doc.settings._element
        # check if updateFields exists
        updateFields = settings.find(qn('w:updateFields'))
        if updateFields is None:
            updateFields = OxmlElement('w:updateFields')
            updateFields.set(qn('w:val'), 'true')
            settings.append(updateFields)
            print("Set updateFields to true in settings.")
    except Exception as e:
        print(f"Could not update settings: {e}")

    doc.save(output_path)
    print(f"Saved to {output_path}")

add_native_toc(r'd:\RESEARCH PNGS\zukhruftheisis_v3.docx', r'd:\RESEARCH PNGS\zukhruftheisis_v4.docx')

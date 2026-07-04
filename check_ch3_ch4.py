from docx import Document
import re

doc = Document("Huma Thesis_Updated.docx")

with open("ch3_ch4_dump.txt", "w", encoding="utf-8") as f:
    f.write("--- Chapter 3 Headings & Suspicious Text ---\n")
    in_ch3 = False
    in_ch4 = False
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        style = para.style.name
        
        if text.upper() == "CHAPTER III":
            in_ch3 = True
        elif text.upper() == "CHAPTER IV":
            in_ch3 = False
            in_ch4 = True
        elif text.upper() == "CHAPTER V":
            in_ch4 = False
            
        is_heading_style = style.startswith("Heading")
        looks_like_heading = re.match(r'^((Chapter\s+\d+)|([1-9](\.\d+)+))', text, re.IGNORECASE) or (len(text) < 100 and text.isupper())
        
        if (in_ch3 or in_ch4) and (is_heading_style or looks_like_heading):
            ch_str = "CH3" if in_ch3 else "CH4"
            f.write(f"[{ch_str} | Para {i} | Style: {style}] {text}\n")
            
    # specifically look at paragraphs 290 to 350 (End of Ch 3)
    f.write("\n--- End of Chapter 3 Text Dump ---\n")
    for i in range(290, 350):
        if i < len(doc.paragraphs):
            text = doc.paragraphs[i].text.strip()
            if text:
                f.write(f"[{i} | {doc.paragraphs[i].style.name}] {text[:100]}...\n")

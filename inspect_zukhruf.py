"""
Inspect zukhruftheisis_revised.docx:
- Find section breaks and page numbering setup
- Find all references in bibliography
- Find all table captions (text under tables)
"""
from docx import Document
from docx.oxml.ns import qn
import re

doc = Document(r'd:\RESEARCH PNGS\zukhruftheisis_revised.docx')

print("=== SECTION BREAKS / PAGE BREAKS IN DOCUMENT ===\n")
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    # Check for page breaks in runs
    for run in para.runs:
        for br in run._r.findall(qn('w:br')):
            btype = br.get(qn('w:type'), '')
            print(f"  [{i}] Break type='{btype}' in para: '{text[:60]}'")
    # Also check sectPr
    pPr = para._p.find(qn('w:pPr'))
    if pPr is not None:
        sectPr = pPr.find(qn('w:sectPr'))
        if sectPr is not None:
            print(f"  [{i}] SECTION BREAK in para: '{text[:80]}'")
            pgNumType = sectPr.find(qn('w:pgNumType'))
            if pgNumType is not None:
                fmt = pgNumType.get(qn('w:fmt'), 'decimal')
                start = pgNumType.get(qn('w:start'), 'not set')
                print(f"       pgNumType: fmt={fmt}, start={start}")

# Check body sectPr
body = doc.element.body
sectPr = body.find(qn('w:sectPr'))
if sectPr is not None:
    print(f"\nFinal body sectPr found:")
    pgNumType = sectPr.find(qn('w:pgNumType'))
    if pgNumType is not None:
        print(f"  pgNumType: fmt={pgNumType.get(qn('w:fmt'))}, start={pgNumType.get(qn('w:start'))}")
    else:
        print("  No pgNumType element")

print("\n\n=== INTRODUCTION START (to find correct paragraph) ===\n")
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if 'CHAPTER I' in text or 'INTRODUCTION' in text.upper():
        print(f"[{i}] Style='{para.style.name}' | '{text[:80]}'")

print("\n\n=== TABLE CAPTIONS (text starting with 'Table' or 'Figure') ===\n")
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    style = para.style.name
    if re.match(r'^(Table|Figure)\s+\d', text, re.IGNORECASE):
        runs = para.runs
        fs = runs[0].font.size if runs else None
        fs_pt = round(fs / 12700) if fs else None
        print(f"[{i}] Style='{style}' | {fs_pt}pt | '{text[:100]}'")

print("\n\n=== REFERENCES SECTION ===\n")
in_refs = False
ref_count = 0
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if 'REFERENCES' in text.upper() and not in_refs:
        in_refs = True
        print(f"References section starts at [{i}]: '{text}'")
        continue
    if in_refs and text:
        ref_count += 1
        print(f"  [{i}] REF {ref_count}: '{text[:120]}'")
print(f"\nTotal reference entries: {ref_count}")

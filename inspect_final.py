"""
Deep inspection of zukhruftheisis_final.docx to understand:
1. Existing TOC entries (if any)
2. Existing footer content (page number fields)
3. Existing section breaks
4. Paragraph structure around TOC area
"""
from docx import Document
from docx.oxml.ns import qn
import re

doc = Document(r'd:\RESEARCH PNGS\zukhruftheisis_final.docx')

print("=== FIRST 60 PARAGRAPHS (structure check) ===\n")
for i, para in enumerate(doc.paragraphs[:60]):
    text = para.text.strip()
    style = para.style.name
    # Check for TOC field
    toc_field = False
    for instr in para._p.findall('.//' + qn('w:instrText')):
        if instr.text and 'TOC' in instr.text:
            toc_field = True
            print(f"[{i}] *** TOC FIELD *** Style='{style}' instrText='{instr.text}'")
    if not toc_field:
        print(f"[{i}] Style='{style}' | '{text[:90]}'")
    # Check for section breaks in pPr
    pPr = para._p.find(qn('w:pPr'))
    if pPr is not None:
        sectPr = pPr.find(qn('w:sectPr'))
        if sectPr is not None:
            pnt = sectPr.find(qn('w:pgNumType'))
            print(f"     ^^ SECTION BREAK HERE: fmt={pnt.get(qn('w:fmt')) if pnt is not None else 'none'}, start={pnt.get(qn('w:start')) if pnt is not None else 'none'}")

print("\n\n=== SECTIONS AND THEIR FOOTERS ===\n")
for si, section in enumerate(doc.sections):
    sectPr = section._sectPr
    pnt = sectPr.find(qn('w:pgNumType'))
    ftr_refs = sectPr.findall(qn('w:footerReference'))
    hdr_refs = sectPr.findall(qn('w:headerReference'))
    print(f"Section {si}:")
    print(f"  pgNumType: {pnt.get(qn('w:fmt')) if pnt is not None else 'none'}, start={pnt.get(qn('w:start')) if pnt is not None else 'none'}")
    print(f"  footerReference count: {len(ftr_refs)}")
    for fr in ftr_refs:
        ftype = fr.get(qn('w:type'))
        rid = fr.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
        print(f"    footer type={ftype}, rId={rid}")
    print(f"  headerReference count: {len(hdr_refs)}")
    
    # Show footer content
    ftr = section.footer
    if ftr:
        for pi, fp in enumerate(ftr.paragraphs):
            full_text = ''.join(t.text or '' for t in fp._p.findall('.//' + qn('w:t')))
            instrs = [i.text for i in fp._p.findall('.//' + qn('w:instrText'))]
            fld_chars = [f.get(qn('w:fldCharType')) for f in fp._p.findall('.//' + qn('w:fldChar'))]
            print(f"  Footer para {pi}: text='{full_text[:60]}' instrs={instrs} fldChars={fld_chars}")
    print()

print("\n\n=== BODY sectPr ===\n")
body = doc.element.body
body_sectPr = body.find(qn('w:sectPr'))
if body_sectPr is not None:
    pnt = body_sectPr.find(qn('w:pgNumType'))
    print(f"Body pgNumType: fmt={pnt.get(qn('w:fmt')) if pnt is not None else 'none'}, start={pnt.get(qn('w:start')) if pnt is not None else 'none'}")
    ftr_refs = body_sectPr.findall(qn('w:footerReference'))
    print(f"Body footer refs: {len(ftr_refs)}")
    for fr in ftr_refs:
        print(f"  type={fr.get(qn('w:type'))}, rId={fr.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')}")
